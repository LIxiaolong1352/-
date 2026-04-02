#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""读取docx文件并提取完整文本内容"""

from docx import Document
import sys
import io

# 设置UTF-8输出
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def read_docx(file_path):
    """读取docx文件并提取文本"""
    doc = Document(file_path)
    
    full_text = []
    total_chars = 0
    
    # 遍历所有段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
            total_chars += len(text)
    
    # 遍历所有表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    full_text.append(text)
                    total_chars += len(text)
    
    return full_text, total_chars

def check_sections(text_content):
    """检查文档是否包含中英文摘要、图表、致谢"""
    
    # 检查中文摘要
    has_cn_abstract = any(keyword in text_content for keyword in ['摘要', '中文摘要'])
    
    # 检查英文摘要
    has_en_abstract = any(keyword in text_content for keyword in ['abstract', 'ABSTRACT'])
    
    # 检查图表
    has_figures = any(keyword in text_content for keyword in ['图', '表', 'Figure', 'Table', '图1', '表1', '图2', '表2'])
    
    # 检查致谢
    has_acknowledgment = any(keyword in text_content for keyword in ['致谢', '致 谢', 'Acknowledgment', 'Acknowledgements'])
    
    return {
        '中文摘要': has_cn_abstract,
        '英文摘要': has_en_abstract,
        '图表': has_figures,
        '致谢': has_acknowledgment
    }

if __name__ == '__main__':
    file_path = r'C:\Users\Administrator\.openclaw\workspace\毕业论文_刘小龙_修改后.docx'
    
    print("=" * 60)
    print("DOCX 文档分析报告")
    print("=" * 60)
    print(f"文件路径: {file_path}")
    print("-" * 60)
    
    # 读取文档
    full_text, total_chars = read_docx(file_path)
    
    # 1. 字数统计
    print(f"\n【1. 字数统计】")
    print(f"总字数（含标点）: {total_chars} 字")
    
    # 合并所有文本用于检查
    all_content = '\n'.join(full_text)
    
    # 2. 检查关键部分
    print(f"\n【2. 文档结构检查】")
    sections = check_sections(all_content)
    for section, exists in sections.items():
        status = "[包含]" if exists else "[未找到]"
        print(f"  {section}: {status}")
    
    # 3. 完整文本内容
    print(f"\n【3. 完整文本内容】")
    print("-" * 60)
    for para in full_text:
        print(para)
    print("-" * 60)
    print(f"\n文档读取完成。共 {len(full_text)} 个段落/单元格。")
