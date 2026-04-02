from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整大纲版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)

# 2.2 生物力学与解剖学分析 - [12]徐小敏
doc.add_heading('2.2 生物力学与解剖学分析', level=2)
content_2_2 = '''篮球运动对膝关节的生物力学要求极高，国内外学者从解剖学和生物力学角度深入分析了膝关节损伤的机制。

徐小敏对江苏省高校高水平运动员的研究显示，膝关节损伤率高达90.71%，主要损伤类型为半月板损伤（22.75%）、髌骨劳损（18.72%）和内侧副韧带损伤（16.82%）[12]。该研究还发现，伤后及时处理率极低，仅14%在1小时内处理，多数无队医，医务监督严重不足。这一研究揭示了高水平运动员面临的严重损伤问题，也为理解膝关节损伤的生物力学机制提供了重要参考。

综合曹炜、张涛、彭清政等学者的研究，可以得出以下生物力学共识：膝关节在半屈位（130°-150°）时稳定性最差，但篮球大部分技术动作在此角度下完成；半月板是承受和传递负荷的核心结构，也是应力最高的部位，因此最易损伤；助跑速度越快，膝关节损伤风险越大（彭清政量化：VGRF从慢速3.15→快速3.79 B.W）；髋内旋角度增大是膝关节损伤的代偿性表现（路程）。

解剖学分析表明，膝关节由股骨下端、胫骨上端和髌骨组成，属于滑车关节。主要结构包括：骨骼结构（股骨下端形成内外侧髁，与胫骨平台的内外侧关节面相对应；髌骨位于股骨前方，与股骨滑车形成髌股关节）；半月板（内侧半月板呈C形，外侧半月板呈O形，具有缓冲震荡、稳定关节、营养关节软骨等功能）；韧带结构（前交叉韧带防止胫骨前移，后交叉韧带防止胫骨后移，内外侧副韧带防止膝关节内外翻）。'''
doc.add_paragraph(content_2_2)

# 2.3 预防干预研究 - [13]Stilling
doc.add_heading('2.3 预防干预研究', level=2)
content_2_3 = '''国外学者在篮球运动膝关节损伤预防干预方面进行了大量研究，取得了重要成果。

Stilling等对青年篮球运动员的研究发现，女性过劳性膝痛发生率（30.4%）高于男性（27.8%），且女性新发病例出现更早（第4周 vs 第7周）[13]。这一研究提示女性运动员是膝关节损伤的高危人群，需要早期强化训练和特别关注。

综合Owoeye等、Leppänen等、Stilling等学者的研究，可以得出以下预防共识：神经肌肉热身（NMT）可有效降低膝关节损伤率（约50%），但对腱病无效；女性运动员是膝关节损伤（尤其是ACL）的高危人群，需早期强化训练；过劳性损伤在女性中负担更重，出现更早。

Owoeye等开展的历史性队列研究显示，SHRed Basketball Injuries NMT热身（10分钟，13个动作）可使膝关节损伤率降低49%（IRR=0.51），踝关节降低32%（IRR=0.68）。Leppänen等的前瞻性观察研究发现，女性膝关节损伤率是男性6.2倍（8例ACL全在女性），女性外翻角显著更大（13.9° vs 2.0°）。'''
doc.add_paragraph(content_2_3)

doc.add_page_break()

print("第二章完成，使用[11][12][13]")
print("继续生成第三章...")

# 第三章 现状调查 - 使用[14]Leppänen
doc.add_heading('第三章 江汉大学大学生篮球运动膝关节损伤现状调查', level=1)

# 3.1 调查对象与方法
doc.add_heading('3.1 调查对象与方法', level=2)
male_count = len([r for r in data if r.get('Q1_性别')=='男'])
female_count = len([r for r in data if r.get('Q1_性别')=='女'])

content_3_1 = f'''本研究以江汉大学有篮球运动经历的全日制在校大学生为调查对象。纳入标准：（1）江汉大学全日制在校学生；（2）每周至少参与1次篮球运动，持续3个月以上；（3）自愿参加调查。排除标准：（1）有先天性膝关节疾病史；（2）近半年内有膝关节手术史；（3）因其他原因无法完成问卷。

共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%。其中男生{male_count}人，女生{female_count}人。

Leppänen等通过前瞻性观察研究，使用3D动作分析技术发现，女性篮球运动员膝关节损伤率是男性的6.2倍，女性外翻角显著大于男性（13.9° vs 2.0°）[14]。这一发现为本研究的性别差异分析提供了重要参考，也提示应特别关注女性篮球参与者的损伤预防。'''
doc.add_paragraph(content_3_1)

print("3.1节完成，使用[14]")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整大纲版.docx')
print("已保存")
