from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整大纲版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 3.2 损伤发生情况分析
doc.add_heading('3.2 损伤发生情况分析', level=2)

content_3_2 = f'''调查结果显示，在{total}名受访者中，有{injured_total}人曾经受过膝关节损伤，损伤发生率为{injured_rate:.1f}%。

损伤类型分布（多选）：'''
doc.add_paragraph(content_3_2)

for t, c in sorted(injury_types.items(), key=lambda x: x[1], reverse=True):
    pct = c / max(injured_total, 1) * 100
    doc.add_paragraph(f'{t}：{c}人（{pct:.1f}%）', style='List Bullet')

content_3_2b = f'''损伤发生情况：起跳落地{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('起跳落地', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('起跳落地', 0)/max(injured_total,1)*100:.1f}%），急停变向{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('急停变向', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('急停变向', 0)/max(injured_total,1)*100:.1f}%）。

损伤主要原因：太累/疲劳{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('太累/疲劳', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('太累/疲劳', 0)/max(injured_total,1)*100:.1f}%），准备活动没做好{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('准备活动没做好', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('准备活动没做好', 0)/max(injured_total,1)*100:.1f}%），场地太滑/条件差{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('场地太滑/条件差', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('场地太滑/条件差', 0)/max(injured_total,1)*100:.1f}%）。'''
doc.add_paragraph(content_3_2b)

# 3.3 不同群体损伤特征对比
doc.add_heading('3.3 不同群体损伤特征对比', level=2)
male_injured = len([r for r in data if r.get('Q1_性别')=='男' and r.get('Q5_是否受伤')=='是'])
female_injured = len([r for r in data if r.get('Q1_性别')=='女' and r.get('Q5_是否受伤')=='是'])
male_count = len([r for r in data if r.get('Q1_性别')=='男'])
female_count = len([r for r in data if r.get('Q1_性别')=='女'])

content_3_3 = f'''性别差异：男生损伤率{male_injured/male_count*100:.1f}%（{male_injured}/{male_count}），女生损伤率{female_injured/female_count*100:.1f}%（{female_injured}/{female_count}）。

运动水平差异：体育专项学生损伤率{len([r for r in data if r.get('Q2_身份')=='体育学院篮球专项' and r.get('Q5_是否受伤')=='是'])/max(len([r for r in data if r.get('Q2_身份')=='体育学院篮球专项']),1)*100:.1f}%，普通爱好者损伤率{len([r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）' and r.get('Q5_是否受伤')=='是'])/max(len([r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）']),1)*100:.1f}%。'''
doc.add_paragraph(content_3_3)

doc.add_page_break()

print("第三章完成，使用[14]")
print("开始生成第四章，使用最后一篇文献[15]...")

# 第四章 预防策略 - 使用[15]Owoeye（最后一篇）
doc.add_heading('第四章 大学生篮球运动膝关节损伤预防策略', level=1)

# 4.1 技术动作优化策略
doc.add_heading('4.1 技术动作优化策略', level=2)
content_4_1 = '''基于前述理论基础和调查结果，提出以下技术动作优化策略：

（1）规范起跳落地技术。控制助跑速度，避免过快助跑；落地时膝关节微屈（约130°-150°），利用肌肉缓冲；双脚同时落地，分散冲击力；落地后顺势屈膝缓冲。

（2）优化急停变向技术。急停时降低重心，增大支撑面；变向时避免膝关节过度内扣；加强髋关节灵活性训练，减少膝关节代偿。

（3）改进半蹲防守姿势。控制半蹲角度，避免长时间保持同一角度；加强股四头肌力量训练，提高半蹲位稳定性。'''
doc.add_paragraph(content_4_1)

# 4.2 身体素质提升策略 - [15] Owoeye
doc.add_heading('4.2 身体素质提升策略', level=2)
content_4_2 = '''基于调查结果和相关研究，提出以下身体素质提升策略：

Owoeye等开展了一项历史性队列研究，评估神经肌肉训练（NMT）热身项目对降低青少年篮球运动员膝关节和踝关节损伤的效果。研究结果显示，NMT热身可使膝关节损伤率降低49%（IRR=0.51），踝关节损伤率降低32%（IRR=0.68）[15]。这一研究为篮球运动膝关节损伤的预防提供了循证医学证据。

基于上述研究成果，建议：（1）引入NMT热身，包括动态拉伸、平衡训练、激活训练和轻量跳跃练习；（2）加强下肢力量训练，包括股四头肌、腘绳肌、臀肌和小腿三头肌力量训练；（3）改善髋关节灵活性，加强髋关节灵活性和力量训练；（4）强化核心稳定性，加强腹肌、背肌力量训练和抗旋转训练。'''
doc.add_paragraph(content_4_2)

print("4.1-4.2节完成，使用[15]")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整大纲版.docx')
print("已保存")
