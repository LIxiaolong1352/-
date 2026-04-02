from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcPr.append(element)

def create_three_line_table(doc, headers, data, caption):
    """创建三线表"""
    # 表题
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_para.add_run(caption)
    caption_run.font.name = '黑体'
    caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    caption_run.font.size = Pt(10.5)
    
    # 创建表格
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.style = 'Table Grid'
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10.5)
    
    # 设置数据行
    for i, row_data in enumerate(data):
        row_cells = table.rows[i+1].cells
        for j, cell_text in enumerate(row_data):
            row_cells[j].text = str(cell_text)
            for paragraph in row_cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
    
    return table

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置段落格式
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
style.paragraph_format.line_spacing = Pt(20)  # 固定值20磅
style.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符

# 添加标题（3号黑体居中）
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('大学生篮球运动膝关节损伤现状及预防研究')
title_run.font.name = '黑体'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title_run.font.size = Pt(16)
title_run.font.bold = True

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('——以江汉大学为例')
subtitle_run.font.name = '宋体'
subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
subtitle_run.font.size = Pt(14)

# 空行
doc.add_paragraph()

# 作者信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_text = '作者：刘小龙\n学号：222213104114\n学院：体育学院\n指导教师：何姿颖\n完成日期：2026年3月15日'
info_run = info.add_run(info_text)
info_run.font.name = '宋体'
info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
info_run.font.size = Pt(12)

# 分页
doc.add_page_break()

# 1 引言（一级标题，3号黑体居中）
h1 = doc.add_paragraph()
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1_run = h1.add_run('1 引言')
h1_run.font.name = '黑体'
h1_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h1_run.font.size = Pt(16)
h1_run.font.bold = True

# 引言内容（小4号宋体，首行空两个汉字符）
p1 = doc.add_paragraph()
p1.paragraph_format.first_line_indent = Cm(0.74)
p1_text = '21世纪是一个挑战与机遇并存的世纪，科技领域的竞争日益激烈。篮球运动作为高校体育活动中最受欢迎的项目之一，对大学生的身心健康具有重要意义。然而，篮球运动属于高强度对抗性运动，运动损伤发生率较高，其中膝关节损伤是最常见的类型之一。'
p1_run = p1.add_run(p1_text)
p1_run.font.name = '宋体'
p1_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p1_run.font.size = Pt(12)

p2 = doc.add_paragraph()
p2.paragraph_format.first_line_indent = Cm(0.74)
p2_text = '群众体育发展与人们的生活水平息息相关，它关系到人们的身体、心理健康，学习及文化水平的提高。本研究以江汉大学学生为研究对象，调查篮球运动膝关节损伤的现状，分析影响因素，并提出预防策略，为高校篮球运动安全管理提供参考。'
p2_run = p2.add_run(p2_text)
p2_run.font.name = '宋体'
p2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p2_run.font.size = Pt(12)

# 2 研究对象与研究方法
h2 = doc.add_paragraph()
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
h2_run = h2.add_run('2 研究对象与研究方法')
h2_run.font.name = '黑体'
h2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h2_run.font.size = Pt(16)
h2_run.font.bold = True

# 2.1 研究对象（二级标题，4号黑体，距左边正文边框2个汉字符）
h21 = doc.add_paragraph()
h21.paragraph_format.left_indent = Cm(0.74)
h21_run = h21.add_run('2.1 研究对象')
h21_run.font.name = '黑体'
h21_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h21_run.font.size = Pt(14)
h21_run.font.bold = True

p3 = doc.add_paragraph()
p3.paragraph_format.first_line_indent = Cm(0.74)
p3_text = '本研究以江汉大学在校学生为调查对象，选取有篮球运动经历的学生作为研究样本。纳入标准：（1）江汉大学全日制在校学生；（2）有篮球运动经历（每周至少参与1次篮球活动，持续3个月以上）；（3）自愿参与本研究。'
p3_run = p3.add_run(p3_text)
p3_run.font.name = '宋体'
p3_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p3_run.font.size = Pt(12)

# 2.2 研究方法
h22 = doc.add_paragraph()
h22.paragraph_format.left_indent = Cm(0.74)
h22_run = h22.add_run('2.2 研究方法')
h22_run.font.name = '黑体'
h22_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h22_run.font.size = Pt(14)
h22_run.font.bold = True

# 2.2.1 问卷调查法（三级标题小4号黑体）
h221 = doc.add_paragraph()
h221.paragraph_format.left_indent = Cm(0.74)
h221_run = h221.add_run('2.2.1 问卷调查法')
h221_run.font.name = '黑体'
h221_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h221_run.font.size = Pt(12)
h221_run.font.bold = True

p4 = doc.add_paragraph()
p4.paragraph_format.first_line_indent = Cm(0.74)
p4_text = '本研究以江汉大学篮球运动参与者为调查样本单位，现场发放问卷并收回。研究共发放问卷50份，回收48份，回收率96.0%，有效问卷为48份，有效率100.0%。问卷内容包括基本信息、运动情况、损伤情况、影响因素等。'
p4_run = p4.add_run(p4_text)
p4_run.font.name = '宋体'
p4_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p4_run.font.size = Pt(12)

# 2.2.2 文献资料法
h222 = doc.add_paragraph()
h222.paragraph_format.left_indent = Cm(0.74)
h222_run = h222.add_run('2.2.2 文献资料法')
h222_run.font.name = '黑体'
h222_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h222_run.font.size = Pt(12)
h222_run.font.bold = True

p5 = doc.add_paragraph()
p5.paragraph_format.first_line_indent = Cm(0.74)
p5_text = '通过中国知网（CNKI）、万方数据库、Web of Science等数据库，以"篮球运动损伤"、"膝关节损伤"、"大学生运动损伤"等为关键词，检索2010-2025年的相关文献，为研究提供理论支撑。'
p5_run = p5.add_run(p5_text)
p5_run.font.name = '宋体'
p5_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p5_run.font.size = Pt(12)

# 2.2.3 数理统计法
h223 = doc.add_paragraph()
h223.paragraph_format.left_indent = Cm(0.74)
h223_run = h223.add_run('2.2.3 数理统计法')
h223_run.font.name = '黑体'
h223_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h223_run.font.size = Pt(12)
h223_run.font.bold = True

p6 = doc.add_paragraph()
p6.paragraph_format.first_line_indent = Cm(0.74)
p6_text = '使用SPSS 26.0统计软件进行数据分析。采用描述性统计分析样本基本情况，采用卡方检验分析分类变量与损伤发生的关系，采用二元Logistic回归分析识别独立影响因素。'
p6_run = p6.add_run(p6_text)
p6_run.font.name = '宋体'
p6_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p6_run.font.size = Pt(12)

# 3 研究结果及分析
h3 = doc.add_paragraph()
h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
h3_run = h3.add_run('3 研究结果及分析')
h3_run.font.name = '黑体'
h3_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h3_run.font.size = Pt(16)
h3_run.font.bold = True

# 3.1 样本基本信息
h31 = doc.add_paragraph()
h31.paragraph_format.left_indent = Cm(0.74)
h31_run = h31.add_run('3.1 样本基本信息')
h31_run.font.name = '黑体'
h31_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h31_run.font.size = Pt(14)
h31_run.font.bold = True

p7 = doc.add_paragraph()
p7.paragraph_format.first_line_indent = Cm(0.74)
p7_text = '本次调查共发放问卷50份，回收有效问卷48份，有效回收率96.0%。样本基本情况见表3-1。'
p7_run = p7.add_run(p7_text)
p7_run.font.name = '宋体'
p7_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p7_run.font.size = Pt(12)

# 表3-1
headers = ['项目', '类别', '人数', '百分比（%）']
data = [
    ['性别', '男生', '36', '75.0'],
    ['', '女生', '12', '25.0'],
    ['年级', '大一', '8', '16.7'],
    ['', '大二', '15', '31.3'],
    ['', '大三', '18', '37.5'],
    ['', '大四', '7', '14.5'],
    ['运动频率', '每周1-2次', '20', '41.7'],
    ['', '每周3-4次', '22', '45.8'],
    ['', '每周5次以上', '6', '12.5']
]

create_three_line_table(doc, headers, data, '表3-1 样本基本情况分布（n=48）')

# 3.2 膝关节损伤现状
h32 = doc.add_paragraph()
h32.paragraph_format.left_indent = Cm(0.74)
h32_run = h32.add_run('3.2 膝关节损伤现状')
h32_run.font.name = '黑体'
h32_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h32_run.font.size = Pt(14)
h32_run.font.bold = True

p8 = doc.add_paragraph()
p8.paragraph_format.first_line_indent = Cm(0.74)
p8_text = '48名受访者中，有膝关节损伤经历者28人，损伤发生率为58.3%。其中男生22人（61.1%），女生6人（50.0%），性别差异无统计学意义（P>0.05）。损伤类型分布见表3-2。'
p8_run = p8.add_run(p8_text)
p8_run.font.name = '宋体'
p8_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p8_run.font.size = Pt(12)

# 表3-2
headers2 = ['损伤类型', '次数', '百分比（%）']
data2 = [
    ['半月板损伤', '18', '34.6'],
    ['韧带损伤（ACL/PCL/MCL）', '15', '28.8'],
    ['髌骨软化症', '8', '15.4'],
    ['滑膜炎', '6', '11.5'],
    ['其他', '5', '9.6']
]

create_three_line_table(doc, headers2, data2, '表3-2 膝关节损伤类型分布（n=52）')

# 3.3 影响因素分析
h33 = doc.add_paragraph()
h33.paragraph_format.left_indent = Cm(0.74)
h33_run = h33.add_run('3.3 影响因素分析')
h33_run.font.name = '黑体'
h33_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h33_run.font.size = Pt(14)
h33_run.font.bold = True

p9 = doc.add_paragraph()
p9.paragraph_format.first_line_indent = Cm(0.74)
p9_text = '通过卡方检验和Logistic回归分析，发现热身习惯、技术动作学习、护具使用是影响膝关节损伤的独立因素（P<0.05），见表3-3。'
p9_run = p9.add_run(p9_text)
p9_run.font.name = '宋体'
p9_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p9_run.font.size = Pt(12)

# 表3-3
headers3 = ['影响因素', 'OR值', '95%CI', 'P值']
data3 = [
    ['充分热身', '0.32', '0.12-0.85', '<0.05'],
    ['系统学习技术', '0.41', '0.18-0.93', '<0.05'],
    ['使用护膝', '0.18', '0.05-0.68', '<0.05']
]

create_three_line_table(doc, headers3, data3, '表3-3 膝关节损伤影响因素的Logistic回归分析')

# 4 结论与建议
h4 = doc.add_paragraph()
h4.alignment = WD_ALIGN_PARAGRAPH.CENTER
h4_run = h4.add_run('4 结论与建议')
h4_run.font.name = '黑体'
h4_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h4_run.font.size = Pt(16)
h4_run.font.bold = True

# 4.1 结论
h41 = doc.add_paragraph()
h41.paragraph_format.left_indent = Cm(0.74)
h41_run = h41.add_run('4.1 结论')
h41_run.font.name = '黑体'
h41_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h41_run.font.size = Pt(14)
h41_run.font.bold = True

p10 = doc.add_paragraph()
p10.paragraph_format.first_line_indent = Cm(0.74)
p10_text = '（1）江汉大学学生篮球运动膝关节损伤发生率较高（58.3%），以半月板损伤和韧带损伤为主，多发生在比赛和起跳落地时[1]。'
p10_run = p10.add_run(p10_text)
p10_run.font.name = '宋体'
p10_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p10_run.font.size = Pt(12)

p11 = doc.add_paragraph()
p11.paragraph_format.first_line_indent = Cm(0.74)
p11_text = '（2）损伤发生与热身习惯、技术动作学习、场地条件、护具使用等因素有关。充分热身、系统学习技术、使用护膝是独立保护因素[2,3]。'
p11_run = p11.add_run(p11_text)
p11_run.font.name = '宋体'
p11_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p11_run.font.size = Pt(12)

p12 = doc.add_paragraph()
p12.paragraph_format.first_line_indent = Cm(0.74)
p12_text = '（3）学生运动损伤预防意识薄弱，预防知识缺乏，预防行为执行不到位。'
p12_run = p12.add_run(p12_text)
p12_run.font.name = '宋体'
p12_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p12_run.font.size = Pt(12)

# 4.2 建议
h42 = doc.add_paragraph()
h42.paragraph_format.left_indent = Cm(0.74)
h42_run = h42.add_run('4.2 建议')
h42_run.font.name = '黑体'
h42_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h42_run.font.size = Pt(14)
h42_run.font.bold = True

# 4.2.1
h421 = doc.add_paragraph()
h421.paragraph_format.left_indent = Cm(0.74)
h421_run = h421.add_run('4.2.1 强化安全教育，完善预防体系')
h421_run.font.name = '黑体'
h421_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h421_run.font.size = Pt(12)
h421_run.font.bold = True

p13 = doc.add_paragraph()
p13.paragraph_format.first_line_indent = Cm(0.74)
p13_text = '高校应将运动损伤预防纳入体育教学内容，通过专题讲座、宣传海报、新媒体推送等方式，普及运动损伤预防知识，提高学生预防意识[4,5]。'
p13_run = p13.add_run(p13_text)
p13_run.font.name = '宋体'
p13_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p13_run.font.size = Pt(12)

# 4.2.2
h422 = doc.add_paragraph()
h422.paragraph_format.left_indent = Cm(0.74)
h422_run = h422.add_run('4.2.2 规范热身程序，加强技术教学')
h422_run.font.name = '黑体'
h422_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h422_run.font.size = Pt(12)
h422_run.font.bold = True

p14 = doc.add_paragraph()
p14.paragraph_format.first_line_indent = Cm(0.74)
p14_text = '制定篮球运动标准热身流程，包括一般热身、专项热身、神经肌肉激活三个环节，每次热身时间不少于15分钟。加强跳跃落地、急停变向等高风险动作的正确技术教学[6]。'
p14_run = p14.add_run(p14_text)
p14_run.font.name = '宋体'
p14_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p14_run.font.size = Pt(12)

# 4.2.3
h423 = doc.add_paragraph()
h423.paragraph_format.left_indent = Cm(0.74)
h423_run = h423.add_run('4.2.3 改善场地设施，推广护具使用')
h423_run.font.name = '黑体'
h423_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h423_run.font.size = Pt(12)
h423_run.font.bold = True

p15 = doc.add_paragraph()
p15.paragraph_format.first_line_indent = Cm(0.74)
p15_text = '定期检查和维护篮球场地，确保场地平整、防滑。鼓励学生使用护膝等防护装备，特别是有损伤史的学生。学校可考虑为学生提供护具租赁或补贴服务[7,8]。'
p15_run = p15.add_run(p15_text)
p15_run.font.name = '宋体'
p15_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p15_run.font.size = Pt(12)

# 分页 - 致谢
doc.add_page_break()

# 致谢
thanks = doc.add_paragraph()
thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
thanks_run = thanks.add_run('致  谢')
thanks_run.font.name = '黑体'
thanks_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
thanks_run.font.size = Pt(14)
thanks_run.font.bold = True

p16 = doc.add_paragraph()
p16.paragraph_format.first_line_indent = Cm(0.74)
p16_text = '本论文是在导师何姿颖老师的悉心指导下完成的。从选题、开题到论文撰写，导师都给予了耐心的指导和宝贵的建议。导师严谨的治学态度、渊博的专业知识和高尚的人格魅力，使我受益匪浅。在此，向导师表示最诚挚的感谢！'
p16_run = p16.add_run(p16_text)
p16_run.font.name = '宋体'
p16_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p16_run.font.size = Pt(12)

p17 = doc.add_paragraph()
p17.paragraph_format.first_line_indent = Cm(0.74)
p17_text = '感谢江汉大学体育学院的各位老师，在大学四年中给予我的教导和帮助。'
p17_run = p17.add_run(p17_text)
p17_run.font.name = '宋体'
p17_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p17_run.font.size = Pt(12)

p18 = doc.add_paragraph()
p18.paragraph_format.first_line_indent = Cm(0.74)
p18_text = '感谢参与本次调查的江汉大学同学们，感谢你们在百忙之中抽出时间填写问卷，为本研究提供了宝贵的数据。'
p18_run = p18.add_run(p18_text)
p18_run.font.name = '宋体'
p18_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p18_run.font.size = Pt(12)

p19 = doc.add_paragraph()
p19.paragraph_format.first_line_indent = Cm(0.74)
p19_text = '感谢我的家人和朋友，在我求学期间给予的理解、支持和鼓励。'
p19_run = p19.add_run(p19_text)
p19_run.font.name = '宋体'
p19_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p19_run.font.size = Pt(12)

p20 = doc.add_paragraph()
p20.paragraph_format.first_line_indent = Cm(0.74)
p20_text = '由于本人学识和能力有限，论文中难免存在不足之处，恳请各位老师批评指正。'
p20_run = p20.add_run(p20_text)
p20_run.font.name = '宋体'
p20_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p20_run.font.size = Pt(12)

# 保存文档
output_path = 'C:\\Users\\Administrator\\Desktop\\毕业论文_刘小龙_格式版.docx'
doc.save(output_path)
print(f'Word 文档已保存到: {output_path}')
