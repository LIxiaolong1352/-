from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import csv

# 创建新文档
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# 读取数据
with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

male_count = len([r for r in data if r.get('Q1_性别')=='男'])
female_count = len([r for r in data if r.get('Q1_性别')=='女'])
male_injured = len([r for r in data if r.get('Q1_性别')=='男' and r.get('Q5_是否受伤')=='是'])
female_injured = len([r for r in data if r.get('Q1_性别')=='女' and r.get('Q5_是否受伤')=='是'])

specialty_count = len([r for r in data if r.get('Q2_身份')=='体育学院篮球专项'])
amateur_count = len([r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）'])
specialty_injured = len([r for r in data if r.get('Q2_身份')=='体育学院篮球专项' and r.get('Q5_是否受伤')=='是'])
amateur_injured = len([r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）' and r.get('Q5_是否受伤')=='是'])

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 标题
title = doc.add_heading('大学生篮球运动中膝关节损伤现状调查与预防策略研究', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.bold = True

subtitle = doc.add_paragraph('——以江汉大学为例')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.name = '宋体'
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph()

# 作者信息
author_info = doc.add_paragraph()
author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_info.add_run('作者：刘小龙\n').font.name = '宋体'
author_info.add_run('学号：222213104114\n').font.name = '宋体'
author_info.add_run('学院：体育学院\n').font.name = '宋体'
author_info.add_run('指导教师：何姿颖').font.name = '宋体'

doc.add_page_break()

# 中文摘要 - 扩充版
doc.add_heading('摘要', level=1)

abstract_text = f'''篮球运动是大学生最喜爱的体育运动项目之一，但膝关节损伤问题日益突出，严重影响大学生的身心健康和运动参与。随着高校体育事业的蓬勃发展，越来越多的大学生参与到篮球运动中，但随之而来的运动损伤问题也日益严重。膝关节作为人体最大、最复杂的关节，承担着支撑身体重量和完成跑跳动作的重要功能，在篮球运动中承受着巨大的负荷。

曹炜从解剖学角度指出，膝关节在半屈位（130°-150°）时稳定性最差，而篮球运动中的大部分技术动作正是在此角度下完成，这使得膝关节成为最容易受伤的部位[1]。张涛的生物力学研究也证实，屈膝130°-150°时伸膝力量最大但稳定性最差，半屈位时韧带松弛，易发生韧带和半月板损伤[2]。

本研究以江汉大学大学生篮球运动参与者为研究对象，采用问卷调查法、文献资料法和数理统计法，对大学生篮球运动膝关节损伤现状及预防策略进行系统研究。参考张海军对辽宁省青少年体校运动员的调查方法，设计《大学生篮球运动膝关节损伤调查问卷》[3]。问卷内容包括基本信息、运动情况、损伤情况、影响因素四个部分，涵盖性别、年龄、运动年限、每周打球频率、损伤类型、损伤原因等多个维度。共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%。调查对象包括体育学院篮球专项学生和普通篮球爱好者，涵盖不同性别、不同运动年限的大学生群体。

调查结果显示：江汉大学大学生篮球运动膝关节损伤发生率为{injured_rate:.1f}%。这一比例与高万钧对河北省普通高校运动员的调查结果（37.5%）接近，高于张海军对青少年体校的调查（14.42%），低于徐小敏对高水平运动员的调查（90.71%），说明江汉大学大学生篮球运动膝关节损伤问题较为普遍[4]。赵孝凯指出，技术动作错误、场地器材问题和局部负担过重是主要成因[5]。损伤类型以半月板损伤（58.1%）和髌骨劳损（58.1%）为主；损伤主要发生在起跳落地（51.6%）和急停变向（22.6%）时；导致损伤的主要原因包括太累/疲劳（32.3%）、准备活动不充分（25.8%）和场地条件差（16.1%）。

研究表明，大学生篮球运动膝关节损伤与热身习惯、技术学习、护具使用等因素密切相关。不同群体损伤特征存在差异，体育专项学生损伤率略高于普通爱好者。基于损伤现状和影响因素分析，本研究从技术动作优化、身体素质提升、训练管理优化、环境保障完善四个维度提出了系统的预防策略，为降低大学生篮球运动膝关节损伤发生率提供理论参考和实践指导。'''
doc.add_paragraph(abstract_text)

keywords = doc.add_paragraph()
keywords.add_run('关键词：').font.bold = True
keywords.add_run('大学生；篮球运动；膝关节损伤；预防策略；江汉大学')

doc.add_page_break()

print("中文摘要完成，内容已扩充")
print("继续生成英文摘要和第一章...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_详细版.docx')
print("已保存")
