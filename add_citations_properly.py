import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import re

doc_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整版.docx'
doc = Document(doc_path)

# 第一步：清除所有现有引用
for para in doc.paragraphs:
    text = para.text
    new_text = re.sub(r'\[\d+\]', '', text)
    if new_text != text:
        para.text = new_text

print('已清除所有引用')

# 第二步：按顺序在合适位置添加15个引用
# 格式: (段落关键词, 引用编号, 是否必须在句末)
citations_plan = [
    # [1] 曹炜(2013) - 解剖学分析 - 放在膝关节解剖相关内容
    ('膝盖里面有半月板', '[1]', False),
    
    # [2] 张涛(2010) - 生物力学 - 放在生物力学相关内容
    ('生物力学', '[2]', False),
    
    # [3] 张海军(2011) - 辽宁青少年调查 - 放在损伤率相关内容
    ('损伤率为', '[3]', False),
    
    # [4] 高万钧(2011) - 河北普通高校调查 - 放在普通高校相关内容
    ('普通高校', '[4]', False),
    
    # [5] 赵孝凯 - 成因分析 - 放在损伤原因分析
    ('损伤原因', '[5]', False),
    
    # [6] 刘虎(2016) - 和田师专调查 - 放在技术动作相关内容
    ('技术动作错误', '[6]', False),
    
    # [7] 路程(2020) - 青少年特点 - 放在青少年/年限相关内容
    ('运动年限', '[7]', False),
    
    # [8] 陈思伟(2021) - 预防策略 - 放在预防相关内容
    ('预防策略', '[8]', False),
    
    # [9] 蒋莉 - 高校损伤 - 放在高校相关内容
    ('高校篮球', '[9]', False),
    
    # [10] 彭清政(2024) - 生物力学特征 - 放在急停起跳相关内容
    ('急停起跳', '[10]', False),
    
    # [11] 张冬(2020) - 泉州调查 - 放在准备活动相关内容
    ('准备活动', '[11]', False),
    
    # [12] 徐小敏(2010) - 江苏高水平 - 放在高水平运动员相关内容
    ('高水平', '[12]', False),
    
    # [13] Stilling(2021) - 过劳性膝痛 - 放在女性/疲劳相关内容
    ('女性', '[13]', False),
    
    # [14] Leppanen(2021) - 女性ACL - 放在女性损伤率相关内容
    ('损伤率', '[14]', False),
    
    # [15] Owoeye(2021) - NMT热身 - 放在热身/预防训练相关内容
    ('热身', '[15]', False),
]

added = []
not_found = []

for keyword, citation, at_end in citations_plan:
    found = False
    for para in doc.paragraphs:
        text = para.text
        if keyword in text and citation not in text:
            # 添加引用
            if at_end:
                # 在段落末尾添加
                para.add_run(citation)
            else:
                # 在关键词后添加
                para.add_run(citation)
            added.append(citation)
            found = True
            print(f'已添加 {citation} 到包含"{keyword[:15]}"的段落')
            break
    
    if not found:
        not_found.append((citation, keyword))

print(f'\n成功添加: {len(added)} 个引用')
print(f'未找到: {len(not_found)} 个')

if not_found:
    print('未找到的引用:')
    for c, k in not_found:
        print(f'  {c} - {k}')

# 保存
output_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx'
doc.save(output_path)
print(f'\n保存完成: {output_path}')

# 验证
doc2 = Document(output_path)
all_citations = []
for para in doc2.paragraphs:
    matches = re.findall(r'\[(\d+)\]', para.text)
    all_citations.extend(matches)

print(f'\n验证: 共 {len(all_citations)} 个引用')
print('引用顺序:', all_citations)
print('唯一引用:', len(set(all_citations)))

# 检查是否按顺序
is_ordered = all_citations == sorted(all_citations, key=int)
no_duplicates = len(all_citations) == len(set(all_citations))

if is_ordered and no_duplicates and len(all_citations) == 15:
    print('✓ 引用正确: 按顺序[1]-[15]，无重复')
else:
    print('✗ 引用有问题')
    if not is_ordered:
        print('  - 顺序不对')
    if not no_duplicates:
        print('  - 有重复')
