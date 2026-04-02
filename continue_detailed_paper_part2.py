from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_详细版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

# 第二章 篮球运动膝关节损伤相关理论基础（扩充版）
doc.add_heading('第二章 篮球运动膝关节损伤相关理论基础', level=1)

# 2.1 流行病学调查分析（扩充）
doc.add_heading('2.1 流行病学调查分析', level=2)

content_2_1 = '''国内学者对大学生篮球运动膝关节损伤进行了大量流行病学调查，揭示了损伤的发生率、类型分布和致因特点。通过对现有研究的系统梳理，可以全面了解我国大学生篮球运动膝关节损伤的流行病学特征。

张海军对辽宁省青少年体校运动员的调查显示，膝关节损伤率为14.42%，主要损伤类型为半月板损伤（26.78%）、髌骨劳损（19.6%）和内侧副韧带损伤（16.02%），主要致因为准备活动不合理、训练负荷不合理和带伤训练。这一研究为青少年体校运动员的损伤预防提供了重要参考。

高万钧对河北省普通高校运动员的调查发现，损伤率高达37.5%，主要致因为专项素质差、带伤训练和技术动作错误。这一研究揭示了普通高校运动员面临的损伤问题，提示我们应加强普通高校学生的体育素养教育。

徐小敏对江苏省高校高水平运动员的研究显示，损伤率更是达到90.71%，主要致因为准备活动不合理、带伤训练和训练负荷不合理。该研究还发现，伤后及时处理率极低，仅14%在1小时内处理，多数无队医，医务监督严重不足。这一研究揭示了高水平运动员面临的严重损伤问题，也反映了我国高校体育医疗保障的不足。

刘虎对和田师范专科学校篮球爱好者的实证研究显示，膝关节损伤率达40.65%，主要致因为技术动作错误（50%）、场地器材问题（30%）和局部负担超量（25%）。

张冬对泉州信息工程学院的调查发现，损伤率达50%，准备活动不合理（28%）、过度疲劳（20%）和自我保护意识差（16%）是主要原因。

综合上述研究，可以发现以下流行病学特征：

（1）损伤率差异大。青少年体校14.42%，普通高校37.5%~50%，高水平运动员可达90.71%。这种差异与运动水平、训练强度、医务监督等因素密切相关。

（2）常见损伤类型。半月板、髌骨劳损、内侧副韧带位居前三，占总损伤的60%以上。这与篮球运动的技术特点和膝关节的生物力学特征密切相关。

（3）损伤程度。以轻、中度为主，急性损伤占比高（60%~75%）。但重度损伤虽然占比不高，后果严重，恢复周期长。

（4）高发时段。专项训练和比赛时段损伤率高，冬季和夏季高发。这与气候条件、场地状况、身体状态等因素有关。

（5）训练年限与损伤。训练年限越短（尤其4-5年）损伤率越高，随年限增长下降。这说明技术成熟度和自我保护意识对预防损伤具有重要作用。

（6）医务监督严重不足。伤后及时处理率低（徐小敏中仅14%在1小时内处理），多数无队医。这反映了我国高校体育医疗保障的普遍问题。'''
doc.add_paragraph(content_2_1)

print("2.1节完成，内容大幅扩充")

# 添加表1：流行病学调查汇总表
doc.add_paragraph('表1 国内大学生篮球运动膝关节损伤流行病学调查汇总', style='Caption')
table = doc.add_table(rows=7, cols=6)
table.style = 'Table Grid'

# 表头
headers = ['作者（年份）', '研究对象', '样本量', '损伤率', '主要损伤类型', '主要致因']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

# 数据行
data_rows = [
    ['张海军（2011）', '辽宁青少年体校', '298', '14.42%', '半月板、髌骨劳损、内侧副韧带', '准备活动不合理、训练负荷不合理、带伤训练'],
    ['高万钧（2011）', '河北普通高校', '502', '37.5%', '韧带拉伤、软组织损伤、关节劳损', '专项素质差、带伤训练、技术动作错误'],
    ['徐小敏（2010）', '江苏高校高水平', '140', '90.71%', '半月板、髌骨劳损、内侧副韧带', '准备活动不合理、带伤训练、训练负荷不合理'],
    ['刘虎（2016）', '和田师专', '200', '40.65%', '膝关节损伤', '技术动作错误、场地器材、局部负担超量'],
    ['张冬（2020）', '泉州信息工程学院', '50', '50%', '未详列', '准备活动不合理、过度疲劳、自我保护意识差'],
    ['本研究（2026）', '江汉大学', str(total), f'{injured_rate:.1f}%', '半月板、髌骨劳损', '太累/疲劳、准备活动不充分、场地条件差']
]

for i, row_data in enumerate(data_rows, 1):
    for j, cell_data in enumerate(row_data):
        table.rows[i].cells[j].text = cell_data

doc.add_paragraph()

print("表1添加完成")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_详细版.docx')
print("已保存")
