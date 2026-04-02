#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
在Word文档中添加文献引用标注 - 完整版
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt
import re

# 定义关键词和对应的引用标注
keywords_refs = [
    ("膝盖里面有半月板", "[1]"),
    ("生物力学", "[2]"),
    ("受伤率达到", "[3]"),
    ("准备活动不充分", "[4]"),
    ("疲劳", "[5]"),
    ("某师范类专科院校", "[6]"),
    ("青少年", "[7]"),
    ("预防对策", "[8]"),
    ("高校篮球运动员", "[9]"),
    ("急停起跳", "[10]"),
    ("调查及预防", "[11]"),
    ("高水平", "[12]"),
    ("过劳性", "[13]"),
    ("女性", "[14]"),
    ("神经肌肉", "[15]"),
]

# 打开文档
doc_path = r"C:\Users\Administrator\.openclaw\workspace\毕业论文_刘小龙_完整版.docx"
output_path = r"C:\Users\Administrator\.openclaw\workspace\毕业论文_刘小龙_引用版.docx"

doc = Document(doc_path)

# 跟踪每个引用是否已添加
ref_added = {ref: False for _, ref in keywords_refs}

# 首先扫描所有段落，查看包含哪些关键词
print("=== 扫描文档中的关键词 ===")
for i, para in enumerate(doc.paragraphs):
    para_text = para.text.strip()
    if para_text:
        for keyword, ref in keywords_refs:
            if keyword in para_text:
                print(f"段落 {i}: 包含 '{keyword}' -> {ref}")

print("\n=== 添加引用标注 ===")

# 遍历所有段落并添加引用
for para in doc.paragraphs:
    para_text = para.text.strip()
    
    # 跳过空段落
    if not para_text:
        continue
    
    # 检查每个关键词
    for keyword, ref in keywords_refs:
        # 如果该引用还未添加，并且段落包含关键词
        if not ref_added[ref] and keyword in para_text:
            # 检查段落末尾是否已有引用标注
            if not re.search(r'\[\d+\]$', para_text.rstrip()):
                # 在段落末尾添加引用
                run = para.add_run(ref)
                run.font.size = Pt(10.5)
                run.font.name = 'Times New Roman'
                ref_added[ref] = True
                print(f"已添加 {ref} 到包含 '{keyword}' 的段落")
                break  # 一个段落只添加一个引用

# 统计添加的引用数量
added_count = sum(1 for v in ref_added.values() if v)
missing_refs = [ref for ref, added in ref_added.items() if not added]

# 保存新文档
doc.save(output_path)
print(f"\n文档已保存到: {output_path}")
print(f"共添加 {added_count} 处引用标注")
if missing_refs:
    print(f"未找到的引用: {', '.join(missing_refs)}")
