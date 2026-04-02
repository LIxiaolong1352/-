import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import re

doc_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整版.docx'
doc = Document(doc_path)

# 清除所有现有引用
for para in doc.paragraphs:
    text = para.text
    new_text = re.sub(r'\[\d+\]', '', text)
    if new_text != text:
        para.text = new_text

print('已清除所有引用')

# 按照段落索引从小到大分配引用 [1]-[15]
# 确保引用按顺序出现在文档中
citations_by_para_order = [
    (6, '[1]'),    # 第6段 - 研究背景
    (31, '[2]'),   # 第31段 - 生物力学
    (37, '[3]'),   # 第37段 - 高校篮球
    (41, '[4]'),   # 第41段 - 损伤率
    (45, '[5]'),   # 第45段 - 高水平
    (47, '[6]'),   # 第47段 - 损伤率
    (51, '[7]'),   # 第51段 - 运动年限
    (53, '[8]'),   # 第53段 - 女性
    (59, '[9]'),   # 第59段 - 损伤原因
    (61, '[10]'),  # 第61段 - 准备活动
    (65, '[11]'),  # 第65段 - 技术动作
    (68, '[12]'),  # 第68段 - 预防策略
    (72, '[13]'),  # 第72段 - 解剖
    (74, '[14]'),  # 第74段 - 急停起跳
    (76, '[15]'),  # 第76段 - 热身
]

added = []
for para_idx, citation in citations_by_para_order:
    if para_idx < len(doc.paragraphs):
        para = doc.paragraphs[para_idx]
        para.add_run(citation)
        added.append(citation)
        print(f'已添加 {citation} 到第{para_idx}段')

print(f'\n共添加 {len(added)} 个引用')

# 保存
output_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx'
doc.save(output_path)
print(f'保存完成: {output_path}')

# 验证 - 按文档顺序读取
doc2 = Document(output_path)
all_citations = []
for para in doc2.paragraphs:
    matches = re.findall(r'\[(\d+)\]', para.text)
    all_citations.extend(matches)

print(f'\n验证: 共 {len(all_citations)} 个引用')
print('文档中出现的引用顺序:', all_citations)

# 检查
expected = [str(i) for i in range(1, 16)]
if all_citations == expected:
    print('✓ 引用正确: 按顺序[1]-[15]')
else:
    print('✗ 顺序不对')
    print('期望:', expected)
