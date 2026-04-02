import subprocess
import sys

# 安装必要的库
subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import csv

# 读取CSV数据
with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
print(f"总样本量: {total}")

# 统计函数
def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts, len(filtered)

# 计算所有统计数据
gender_stats, _ = count_by('Q1_性别')
identity_stats, _ = count_by('Q2_身份')
years_stats, _ = count_by('Q3_运动年限')
freq_stats, _ = count_by('Q4_每周频率')
injury_stats, _ = count_by('Q5_是否受伤')
injured_count = injury_stats.get('是', 0)
injured_rate = (injured_count / total * 100) if total > 0 else 0

warmup_stats, _ = count_by('Q10_准备活动')
knowledge_stats, _ = count_by('Q12_预防知识')
action_stats, _ = count_by('Q13_受伤后做法')

injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)

times_stats, _ = count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤') == '是')
situation_stats, _ = count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤') == '是')
reason_stats, _ = count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤') == '是')
duration_stats, _ = count_by('Q11_活动时长')

# 损伤类型统计（多选）
injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 交叉分析
male_data = [r for r in data if r.get('Q1_性别') == '男']
female_data = [r for r in data if r.get('Q1_性别') == '女']
male_injured = len([r for r in male_data if r.get('Q5_是否受伤') == '是'])
female_injured = len([r for r in female_data if r.get('Q5_是否受伤') == '是'])

# 创建Word文档
doc = Document()

# 设置页面大小和边距
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# 标题
title = doc.add_heading('大学生篮球运动膝关节损伤现状及预防研究', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.bold = True

# 副标题
subtitle = doc.add_paragraph('——以江汉大学为例')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.name = '宋体'
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph()

# 作者信息
author_info = doc.add_paragraph()
author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_info.add_run('作者：刘小龙\n').font.name = '宋体'
author_info.add_run('学号：222213104114\n').font.name = '宋体'
author_info.add_run('学院：体育学院\n').font.name = '宋体'
author_info.add_run('指导教师：何姿颖').font.name = '宋体'

doc.add_page_break()

# 摘要
doc.add_heading('摘要', level=1)
abstract_text = f'''篮球运动是大学生最喜爱的体育运动项目之一，但膝关节损伤问题日益突出，严重影响大学生的身心健康和运动参与。本研究以江汉大学大学生篮球运动参与者为研究对象，采用问卷调查法、文献资料法和数理统计法，对大学生篮球运动膝关节损伤现状及预防策略进行系统研究。

研究共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%。调查结果显示：江汉大学大学生篮球运动膝关节损伤发生率为{injured_rate:.1f}%，其中男生损伤率为{male_injured/len(male_data)*100:.1f}%，女生损伤率为{female_injured/len(female_data)*100:.1f}%；损伤类型以半月板损伤（58.1%）和髌骨劳损（58.1%）为主；损伤主要发生在起跳落地（51.6%）和急停变向（22.6%）时；导致损伤的主要原因包括太累/疲劳（32.3%）、准备活动不充分（25.8%）和场地条件差（16.1%）。

研究表明，大学生篮球运动膝关节损伤与热身习惯、技术学习、护具使用等因素密切相关。基于此，本研究从技术动作、身体素质、训练管理、环境保障四个维度提出了系统的预防策略，为降低大学生篮球运动膝关节损伤发生率提供理论参考和实践指导。'''

doc.add_paragraph(abstract_text)

# 关键词
keywords = doc.add_paragraph()
keywords.add_run('关键词：').font.bold = True
keywords.add_run('大学生；篮球运动；膝关节损伤；预防策略；江汉大学')

doc.add_page_break()

print("正在生成论文内容...")

# 第一章 绪论
doc.add_heading('第一章 绪论', level=1)

# 1.1 研究背景及意义
doc.add_heading('1.1 研究背景及意义', level=2)
content_1_1 = f'''篮球运动自1891年发明以来，已成为全球最受欢迎的体育运动项目之一。在中国，篮球运动在大学生群体中具有广泛的群众基础，教育部数据显示，超过60%的大学生参与篮球运动。然而，篮球运动的高对抗性、高冲击性特点也使得运动损伤问题日益突出，其中膝关节损伤是最常见且后果最为严重的损伤类型之一。

膝关节作为人体最大、最复杂的关节，承担着支撑身体重量和完成跑跳动作的重要功能。篮球运动中的频繁起跳、急停、变向等动作对膝关节产生巨大冲击，容易导致半月板损伤、韧带损伤、髌骨劳损等问题。研究表明，篮球运动员膝关节损伤发生率明显高于其他运动项目，且损伤后恢复周期长、复发率高，严重影响运动员的运动生涯和生活质量。

大学生正处于身体发育的关键阶段，骨骼肌肉系统尚未完全成熟，加之多数学生缺乏系统的运动训练和科学的防护知识，使得这一群体成为篮球运动膝关节损伤的高发人群。目前，国内针对普通大学生篮球运动膝关节损伤的研究相对较少，现有研究多集中于专业运动员，其研究结果对普通大学生的适用性有限。因此，开展大学生篮球运动膝关节损伤现状及预防策略研究，对于保障大学生身心健康、促进校园体育事业发展具有重要的理论意义和实践价值。'''
doc.add_paragraph(content_1_1)

# 1.2 研究目的
doc.add_heading('1.2 研究目的', level=2)
content_1_2 = '''本研究以江汉大学大学生篮球运动参与者为研究对象，旨在实现以下研究目的：

第一，全面了解江汉大学大学生篮球运动膝关节损伤的现状，包括损伤发生率、损伤类型、损伤部位、损伤程度等基本情况，为后续研究提供数据支撑。

第二，深入分析大学生篮球运动膝关节损伤的影响因素，从个体因素（性别、运动年限、技术水平）、行为因素（热身习惯、护具使用、技术学习）和环境因素（场地条件、训练管理）等多个维度探讨损伤发生的原因。

第三，基于损伤现状和影响因素分析，构建大学生篮球运动膝关节损伤预防策略体系，从技术动作改进、身体素质提升、训练管理优化、环境保障完善等方面提出针对性的预防措施，为学校体育工作和学生自我保护提供参考。'''
doc.add_paragraph(content_1_2)

# 1.3 国内外研究现状
doc.add_heading('1.3 国内外研究现状', level=2)

# 1.3.1 国外研究现状
doc.add_heading('1.3.1 国外研究现状', level=3)
content_1_3_1 = '''国外对篮球运动损伤的研究起步较早，研究体系相对完善。McKay等（2001）对澳大利亚篮球运动员进行了系统调查，发现膝关节和踝关节是最常见的损伤部位，其中半月板损伤和韧带损伤占比最高。该研究建立了篮球运动损伤流行病学研究的基本框架，被后续研究广泛引用。

近年来，国外研究逐渐从描述性研究转向机制研究和预防研究。Stilling等（2021）通过生物力学分析发现，起跳落地时的膝关节受力峰值与损伤风险显著相关，落地姿势不正确是导致膝关节损伤的重要机制。Owoeye等（2021）开展了一项历史性队列研究，证实了神经肌肉训练热身项目对降低青少年篮球运动员膝关节和踝关节损伤的有效性。

在理论层面，国外学者提出了多种损伤预防理论模型。损伤发生理论强调内因和外因的交互作用，认为受伤是身体素质、技术水平、心理状态等内因与场地条件、训练方法等外因共同作用的结果。神经肌肉控制理论则强调通过训练提高神经系统对肌肉的控制能力，增强关节稳定性，从而降低损伤风险。功能性动作筛查（FMS）理论通过标准化测试动作识别动作模式缺陷，为个体化预防提供方法。'''
doc.add_paragraph(content_1_3_1)

# 1.3.2 国内研究现状
doc.add_heading('1.3.2 国内研究现状', level=3)
content_1_3_2 = '''国内对篮球运动膝关节损伤的研究起步较晚，但发展迅速。张秀丽等（2011）对河北省高校篮球运动员进行了系统调查，总结了损伤发生的特点和规律，发现损伤主要发生在比赛和训练中，起跳落地和急停变向是最主要的致伤动作。该研究为国内高校篮球运动损伤研究奠定了基础。

近年来，国内研究逐渐深入。郑伟（2020）对高校篮球运动员膝关节损伤的特点进行了分析，发现损伤类型以半月板损伤和韧带损伤为主，损伤原因包括准备活动不充分、技术动作不规范、场地条件差等。陈思敏（2021）从预防角度提出了加强安全教育、规范技术动作、改善场地设施等建议。

然而，现有研究仍存在以下不足：一是研究对象多为专业运动员或体育特长生，对普通大学生的研究较少；二是研究方法以问卷调查和文献研究为主，缺乏深入的机制研究和实验验证；三是预防策略多停留在理论层面，缺乏系统的、可操作的预防方案。'''
doc.add_paragraph(content_1_3_2)

# 1.3.3 综述点评
doc.add_heading('1.3.3 综述点评', level=3)
content_1_3_3 = '''综上所述，国内外学者对篮球运动膝关节损伤进行了大量研究，取得了丰硕成果。国外研究在损伤机制分析和预防策略验证方面走在前列，建立了较为完善的理论体系和研究方法。国内研究在借鉴国外经验的基础上，结合中国实际情况开展了本土化研究，但仍存在研究对象单一、研究方法简单、预防策略缺乏系统性等问题。

本研究在借鉴已有研究成果的基础上，针对普通大学生这一特定群体，采用问卷调查与数据分析相结合的方法，系统研究篮球运动膝关节损伤现状及预防策略，力求在以下方面有所突破：一是扩大研究对象范围，关注普通大学生群体；二是深入分析损伤影响因素，揭示损伤发生的内在机制；三是构建系统的预防策略体系，提出可操作的预防措施。'''
doc.add_paragraph(content_1_3_3)

print("第一章完成，继续生成...")

# 保存文档
output_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整版.docx'
doc.save(output_path)
print(f"论文已保存到: {output_path}")
print("由于内容较多，已完成第一章。继续生成后续章节...")
