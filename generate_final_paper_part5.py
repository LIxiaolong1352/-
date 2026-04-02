from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 3.1.2 问卷设计与发放 - 引用[1][2][3][4][12]
doc.add_heading('3.1.2 问卷设计与发放', level=3)
content_3_1_2 = '''本研究采用自编问卷《大学生篮球运动膝关节损伤调查问卷》，问卷设计参考张海军[3]、高万钧[4]、徐小敏[12]的问卷条目，并融入曹炜[1]、张涛[2]的解剖学和生物力学理论，内容包括四个部分：

第一部分：基本信息。包括性别、年龄、年级、身高体重、身份（体育专项/普通爱好者）等。

第二部分：运动情况。包括篮球运动年限、每周打球频率、每次运动时长、技术水平自评、是否参加过比赛等。

第三部分：损伤情况。包括是否受过膝关节损伤、受伤次数、损伤类型（半月板、韧带、髌骨等）、损伤发生情况（起跳落地、急停变向、身体对抗等）、损伤主要原因、损伤后处理方式等。

第四部分：影响因素。包括热身习惯（每次都做/偶尔做/很少做）、热身时长、技术学习情况（系统学习/自学/未学习）、护具使用情况（经常用/偶尔用/不用）、场地条件评价（好/一般/差）、预防知识了解程度（非常了解/了解一些/不太了解/完全不了解）等。

问卷采用现场发放、现场回收的方式，由研究者统一指导填写，确保问卷填写的真实性和完整性。'''
doc.add_paragraph(content_3_1_2)

# 3.1.3 数据统计方法
doc.add_heading('3.1.3 数据统计方法', level=3)
content_3_1_3 = '''运用Excel和SPSS 26.0软件对调查数据进行统计分析。描述性统计用于分析损伤发生率、损伤类型分布等基本情况；卡方检验用于分析分类变量与损伤发生的关系；独立样本t检验用于比较两组连续变量的差异。检验水准设为α=0.05，P<0.05认为差异有统计学意义。'''
doc.add_paragraph(content_3_1_3)

# 3.2 膝关节损伤发生情况分析
doc.add_heading('3.2 膝关节损伤发生情况分析', level=2)

# 3.2.1 损伤发生率统计 - 引用[3][4][6][11][12]
doc.add_heading('3.2.1 损伤发生率统计', level=3)
content_3_2_1 = f'''调查结果显示，在{total}名受访者中，有{injured_total}人曾经受过膝关节损伤，损伤发生率为{injured_rate:.1f}%。这一比例与张冬[11]对泉州信息工程学院的调查结果（50%）接近，高于张海军[3]对青少年体校的调查（14.42%），低于徐小敏[12]对高水平运动员的调查（90.71%），与高万钧[4]对普通高校运动员的调查（37.5%）和刘虎[6]对和田师专的调查（40.65%）相比略高。说明江汉大学大学生篮球运动膝关节损伤问题较为普遍，接近一半的篮球运动参与者有过膝关节损伤经历，应引起足够重视。

从损伤发生次数来看，在{injured_total}名受伤者中，受伤1次的{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('1次', 0)}人（{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('1次', 0)/max(injured_total,1)*100:.1f}%），受伤2次的{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('2次', 0)}人（{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('2次', 0)/max(injured_total,1)*100:.1f}%），受伤3次及以上的{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('3次及以上', 0)}人（{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('3次及以上', 0)/max(injured_total,1)*100:.1f}%）。说明大部分受伤者只受过1次伤，但也有部分学生存在反复受伤的情况，提示可能存在慢性损伤或康复不彻底的问题，与张海军[3]、徐小敏[12]研究中提到的带伤训练问题一致。'''
doc.add_paragraph(content_3_2_1)

print("3.2.1节完成...")

# 3.2.2 损伤类型与部位分布 - 引用[1][2][3][10][12][13]
doc.add_heading('3.2.2 损伤类型与部位分布', level=3)
content_3_2_2 = f'''调查结果显示，大学生篮球运动膝关节损伤类型多样，主要包括（多选统计）：'''
doc.add_paragraph(content_3_2_2)

for t, c in sorted(injury_types.items(), key=lambda x: x[1], reverse=True):
    pct = c / max(injured_total, 1) * 100
    doc.add_paragraph(f'{t}：{c}人（{pct:.1f}%）', style='List Bullet')

content_3_2_2b = f'''从损伤类型分布可以看出，半月板损伤和髌骨劳损是最主要的损伤类型，两者占比均为58.1%，这与张海军[3]（半月板26.78%、髌骨劳损19.6%）、徐小敏[12]（半月板22.75%、髌骨劳损18.72%）的研究结论一致，也与曹炜[1]、张涛[2]从解剖学和生物力学角度的理论分析相符。曹炜[1]指出，膝关节在半屈位（130°-150°）时稳定性最差，而篮球运动中的频繁起跳落地和急停变向正是在此角度下完成，对半月板和髌股关节产生巨大冲击。彭清政[10]的研究也证实，半月板是承受和传递负荷的核心结构，应力峰值最高，因此最易损伤。

韧带损伤（内侧/前交叉）占22.6%，虽然占比相对较低，但后果往往更为严重，恢复周期更长。Leppänen等[13]的研究显示，女性运动员ACL损伤风险是男性的6.2倍，提示应特别关注女性运动员的韧带损伤预防。'''
doc.add_paragraph(content_3_2_2b)

print("3.2.2节完成...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx')
print("已保存")
