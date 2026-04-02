from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 创建新文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

# 添加标题
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run('外文文献译文')
run.font.size = Pt(18)
run.font.bold = True
run.font.name = '黑体'

doc.add_paragraph()

# 添加文章标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.add_run('神经肌肉训练热身方案对减少青少年篮球运动员膝关节和踝关节损伤的有效性：一项历史性队列研究')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = '黑体'

doc.add_paragraph()

# 作者信息
author = doc.add_paragraph('作者：Oluwatoyosi Owoeye, Kati Pasanen, Kimberley Befus, Carlyn Stilling, Brianna Ghali, Tyler J Tait, Tate HubkaRao, Luz Palacios-Derflingher, Vineetha Warriyar, Carolyn Emery')
author.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

doc.add_paragraph('文献编号：10.1136/bjsports-2021-IOC.53')
doc.add_paragraph()

# 正文内容
sections = [
    ('背景', 
     '评估神经肌肉训练（NMT）热身方案在减少青少年篮球运动员膝关节和踝关节损伤方面的有效性的研究较少，特别是NMT热身方案对髌骨和跟腱病变的影响尚不清楚。'),
    
    ('目的', 
     '评估NMT热身方案在降低青少年篮球运动员膝关节和踝关节损伤（包括髌骨和跟腱病变）风险方面的有效性。'),
    
    ('设计', 
     '对暴露于NMT干预（第2赛季）和未暴露（第1赛季）的球员进行两个赛季的历史性队列比较研究。'),
    
    ('研究场所', 
     '青少年篮球队（加拿大阿尔伯塔省）。'),
    
    ('研究对象', 
     '94支球队，共825名男性和女性球员（年龄范围：11-18岁；第1赛季，n=518；第2赛季，n=307）。'),
    
    ('干预措施', 
     '在第2赛季实施由教练指导的10分钟SHRed篮球损伤NMT热身方案，包含13个练习，包括有氧、敏捷、力量和平衡训练。对照组在第1赛季使用他们的标准热身练习。'),
    
    ('主要结果测量', 
     '使用经过验证的损伤监测方法，在两个篮球赛季期间每周记录所有膝关节和踝关节损伤，包括髌骨和跟腱病变。使用泊松回归（使用暴露小时数作为偏移量，并调整团队聚集、性别、两个赛季参与情况）来估计两个赛季之间所有损伤的发病率比率（IRR；98.8%CI Bonferroni校正）。使用逻辑回归（调整团队聚集、性别、暴露小时数、两个赛季参与情况）来估计报告至少一次肌腱病变的球员的优势比（OR；98.8%CI）。'),
    
    ('结果', 
     'NMT热身方案对膝关节损伤[IRR=0.51（98.8%CI：0.35-0.75）]和踝关节损伤[IRR=0.68（98.8%CI：0.52-0.91）]具有保护作用，但对髌骨病变[OR=0.88（98.8%CI：0.44-1.73）]和跟腱病变[OR=0.63（98.8%CI：0.18-2.18）] specifically 无效。'),
    
    ('结论', 
     'SHRed篮球损伤NMT热身方案在降低青少年篮球运动员所有膝关节和踝关节损伤率方面有效，但在减轻髌骨和跟腱病变风险方面无效。进一步评估负荷调整的研究可能是预防肌腱病变的目标。'),
]

for title_text, content in sections:
    # 添加小标题
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.name = '黑体'
    
    # 添加内容
    doc.add_paragraph(content)
    doc.add_paragraph()

# 添加译者注
note = doc.add_paragraph()
run = note.add_run('译者注')
run.font.bold = True
run.font.name = '黑体'

doc.add_paragraph('本文献为Owoeye等学者关于NMT热身对青少年篮球运动员膝关节和踝关节损伤预防效果的研究，发现NMT可使膝关节损伤率降低49%（IRR=0.51），踝关节损伤率降低32%（IRR=0.68）。该研究对理解NMT在青少年篮球损伤预防中的应用具有重要参考价值。')

# 保存文档
output_path = r'C:\Users\Administrator\Desktop\外文文献译文.docx'
doc.save(output_path)

print('译文文档已保存到桌面：外文文献译文.docx')
print('内容包括：')
print('- 标题：神经肌肉训练热身方案对减少青少年篮球运动员膝关节和踝关节损伤的有效性')
print('- 作者信息')
print('- 背景、目的、设计、研究场所、研究对象')
print('- 干预措施、主要结果测量')
print('- 结果：NMT使膝关节损伤降低49%，踝关节降低32%')
print('- 结论')
print('- 译者注')
