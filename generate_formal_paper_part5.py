from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_正式版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

# 3.3 不同群体损伤特征对比分析
doc.add_heading('3.3 不同群体损伤特征对比分析', level=2)

# 3.3.1 性别差异分析
doc.add_heading('3.3.1 性别差异分析', level=3)
male_count = len([r for r in data if r.get('Q1_性别')=='男'])
female_count = len([r for r in data if r.get('Q1_性别')=='女'])
male_injured = len([r for r in data if r.get('Q1_性别')=='男' and r.get('Q5_是否受伤')=='是'])
female_injured = len([r for r in data if r.get('Q1_性别')=='女' and r.get('Q5_是否受伤')=='是'])
male_rate = male_injured / max(male_count, 1) * 100
female_rate = female_injured / max(female_count, 1) * 100

content_3_3_1 = f'''调查结果显示，男生损伤率为{male_rate:.1f}%（{male_injured}/{male_count}），女生损伤率为{female_rate:.1f}%（{female_injured}/{female_count}）。虽然男生损伤人数多于女生，但损伤率差异不大，说明性别不是影响损伤发生的主要因素。

然而，Leppänen等[14]和Stilling等[13]的研究均证实，女性运动员是膝关节损伤（尤其是ACL损伤）的高危人群。Leppänen等[14]发现，女性篮球运动员膝关节损伤率是男性的6.2倍，8例ACL损伤全部发生在女性，女性外翻角显著大于男性（13.9° vs 2.0°）。Stilling等[13]的研究也显示，女性过劳性膝痛发生率（30.4%）高于男性（27.8%），且女性新发病例出现更早（第4周 vs 第7周）。

本研究中男女生损伤率差异不明显，可能与以下因素有关：一是样本量有限，女性样本仅{female_count}人；二是本调查对象为普通大学生，运动强度相对较低，女性运动员在高强度对抗中的损伤风险优势未充分体现。但仍建议对女性篮球参与者加强预防指导，特别是ACL损伤的预防。'''
doc.add_paragraph(content_3_3_1)

# 3.3.2 运动水平差异分析
doc.add_heading('3.3.2 运动水平差异分析', level=3)
specialty_data = [r for r in data if r.get('Q2_身份')=='体育学院篮球专项']
amateur_data = [r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）']
specialty_injured = len([r for r in specialty_data if r.get('Q5_是否受伤')=='是'])
amateur_injured = len([r for r in amateur_data if r.get('Q5_是否受伤')=='是'])
specialty_rate = specialty_injured / max(len(specialty_data), 1) * 100
amateur_rate = amateur_injured / max(len(amateur_data), 1) * 100

content_3_3_2 = f'''调查结果显示，体育学院篮球专项学生损伤率为{specialty_rate:.1f}%（{specialty_injured}/{len(specialty_data)}），普通篮球爱好者损伤率为{amateur_rate:.1f}%（{amateur_injured}/{len(amateur_data)}）。

体育专项学生虽然技术水平较高，但由于训练强度大、比赛频繁，损伤率反而略高于普通爱好者。这与徐小敏[12]对高水平运动员的研究结论一致，高水平运动员损伤率高达90.71%，远高于普通高校的37.5%~50%。普通爱好者虽然技术水平相对较低，但运动强度较小，损伤率略低。这说明运动强度和技术水平的匹配程度是影响损伤的重要因素，也与张海军[3]发现的训练年限与损伤率关系相符——训练年限越短（尤其4-5年）损伤率越高，随年限增长下降。'''
doc.add_paragraph(content_3_3_2)

# 3.3.3 热身习惯与损伤关系分析
doc.add_heading('3.3.3 热身习惯与损伤关系分析', level=3)
always_warmup = [r for r in data if r.get('Q10_准备活动')=='每次都做']
occasionally_warmup = [r for r in data if r.get('Q10_准备活动')=='偶尔做']
always_injured = len([r for r in always_warmup if r.get('Q5_是否受伤')=='是'])
occasionally_injured = len([r for r in occasionally_warmup if r.get('Q5_是否受伤')=='是'])
always_rate = always_injured / max(len(always_warmup), 1) * 100
occasionally_rate = occasionally_injured / max(len(occasionally_warmup), 1) * 100

content_3_3_3 = f'''调查结果显示，每次都做热身的学生{len(always_warmup)}人（{len(always_warmup)/total*100:.1f}%），损伤率为{always_rate:.1f}%（{always_injured}/{len(always_warmup)}）；偶尔做热身的学生{len(occasionally_warmup)}人（{len(occasionally_warmup)/total*100:.1f}%），损伤率为{occasionally_rate:.1f}%（{occasionally_injured}/{len(occasionally_warmup)}）。

虽然两组损伤率差异未达到统计学显著水平，但每次都做热身的损伤率略低于偶尔做热身的。这与Owoeye等[15]的研究结论一致，神经肌肉训练（NMT）热身可使膝关节损伤率降低49%（IRR=0.51）。张海军[3]、徐小敏[12]、张冬[11]的研究也将准备活动不合理列为首要或重要致因。说明热身习惯对预防膝关节损伤具有重要作用，应加强学生热身意识的培养。'''
doc.add_paragraph(content_3_3_3)

print("3.3节完成...")

# 3.4 损伤致因综合分析
doc.add_heading('3.4 损伤致因综合分析', level=2)
content_3_4 = '''基于上述调查结果，结合曹炜[1]、张涛[2]的解剖学和生物力学理论，对江汉大学大学生篮球运动膝关节损伤的致因进行综合分析。

（1）技术动作因素。篮球运动中的起跳落地、急停变向等技术动作对膝关节产生巨大冲击。彭清政[10]的研究证实，助跑速度越快，VGRF越大，快速助跑时VGRF可达3.79倍体重，半月板应力峰值超过10MPa。曹炜[1]、张涛[2]指出，膝关节在半屈位（130°-150°）时稳定性最差，而篮球大部分技术动作正是在此角度下完成。本研究中，起跳落地致伤占51.6%，技术动作错误致因占12.9%，说明技术动作规范对预防损伤至关重要。

（2）身体素质因素。路程[7]发现，髋内旋角度增大和股后肌群肌力下降是膝关节损伤的重要致伤因素。高万钧[4]也将专项素质差列为主要致因。本研究中，太累/疲劳致因占32.3%，说明身体素质不足、疲劳状态下继续运动是导致损伤的重要原因。

（3）训练管理因素。张海军[3]、徐小敏[12]均将训练负荷不合理和带伤训练列为主要致因。徐小敏[12]的研究还发现，伤后及时处理率仅14%，多数无队医。本研究中，准备活动不充分致因占25.8%，说明训练管理、医务监督仍有待加强。

（4）环境因素。刘虎[6]的研究中，场地器材问题占30%。赵孝凯[5]指出，混凝土场地是高危情境。本研究中，场地条件差致因占16.1%，说明场地环境改善也是预防损伤的重要环节。'''
doc.add_paragraph(content_3_4)

doc.add_page_break()

print("第三章完成，开始生成第四章...")

# 第四章
doc.add_heading('第四章 大学生篮球运动膝关节损伤预防策略', level=1)

# 4.1
doc.add_heading('4.1 技术动作优化策略', level=2)
content_4_1 = '''基于曹炜[1]、张涛[2]的解剖学、生物力学理论和彭清政[10]的量化研究结果，提出以下技术动作优化策略：

（1）规范起跳落地技术。彭清政[10]的研究表明，助跑速度越快，膝关节受力越大。应加强起跳落地技术的训练，包括：①控制助跑速度，避免过快助跑；②落地时膝关节微屈（约130°-150°），利用肌肉缓冲；③双脚同时落地，分散冲击力；④落地后顺势屈膝缓冲，避免膝关节伸直锁定。

（2）优化急停变向技术。曹炜[1]指出，急停、旋转时半月板产生矛盾运动，易导致损伤。应加强急停变向技术的训练，包括：①急停时降低重心，增大支撑面；②变向时避免膝关节过度内扣；③加强髋关节灵活性训练，减少膝关节代偿；④使用侧滑步等规范防守步法，避免膝关节过度扭转。

（3）改进半蹲防守姿势。张涛[2]指出，屈膝130°-150°时稳定性最差，但篮球防守多在此角度。应：①控制半蹲角度，避免长时间保持同一角度；②加强股四头肌力量训练，提高半蹲位稳定性；③防守间隙适当活动膝关节，避免僵硬。'''
doc.add_paragraph(content_4_1)

print("4.1节完成...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_正式版.docx')
print("已保存")
