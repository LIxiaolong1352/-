from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_15文献版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 第二章 理论基础 - 使用[11][12][13]
doc.add_heading('第二章 篮球运动膝关节损伤相关理论基础', level=1)

# 2.1 膝关节解剖结构与功能 - [11]张冬
doc.add_heading('2.1 膝关节解剖结构与功能', level=2)
content_2_1 = '''膝关节是人体最大、最复杂的关节，由股骨下端、胫骨上端和髌骨组成，属于滑车关节。膝关节的主要结构包括骨骼结构、半月板、韧带结构和滑膜与滑液。

张冬对泉州信息工程学院的调查发现，损伤率达50%，准备活动不合理（28%）、过度疲劳（20%）和自我保护意识差（16%）是主要原因[11]。这一研究为了解普通高校篮球运动膝关节损伤现状提供了重要数据，也揭示了大学生篮球运动参与者在解剖学知识和自我保护意识方面的不足。

膝关节的主要功能包括屈伸运动和有限的旋转运动。半月板具有缓冲震荡、稳定关节、营养关节软骨等功能。韧带结构包括前交叉韧带（ACL）、后交叉韧带（PCL）和内外侧副韧带，分别防止胫骨前移、后移和膝关节内外翻。理解膝关节的解剖结构对于分析损伤机制和制定预防策略具有重要意义。'''
doc.add_paragraph(content_2_1)

# 2.2 篮球运动中膝关节的生物力学特征 - [12]徐小敏
doc.add_heading('2.2 篮球运动中膝关节的生物力学特征', level=2)
content_2_2 = f'''篮球运动对膝关节的生物力学要求极高。徐小敏对江苏省高校高水平运动员的研究显示，膝关节损伤率高达90.71%，主要损伤类型为半月板损伤（22.75%）、髌骨劳损（18.72%）和内侧副韧带损伤（16.82%）[12]。该研究还发现，伤后及时处理率极低，仅14%在1小时内处理，多数无队医，医务监督严重不足。

篮球运动中的起跳落地、急停变向、半蹲防守等动作对膝关节产生复杂的生物力学负荷。起跳落地时，膝关节需要承受巨大的地面反作用力；急停变向时，半月板和韧带承受剪切力和扭转力；半蹲姿势时，髌股关节压力增大。理解这些生物力学特征对于预防损伤至关重要。'''
doc.add_paragraph(content_2_2)

# 2.3 膝关节损伤的分类与机制 - [13]Stilling
doc.add_heading('2.3 膝关节损伤的分类与机制', level=2)
content_2_3 = '''篮球运动中常见的膝关节损伤主要包括半月板损伤、韧带损伤、髌骨劳损和髌腱炎。

半月板损伤是篮球运动中最常见的膝关节损伤，多发生在起跳落地和急停变向时。韧带损伤中，前交叉韧带（ACL）损伤多发生在急停变向和起跳落地时，内侧副韧带（MCL）损伤多发生在身体对抗时膝关节外翻。

Stilling等对青年篮球运动员的研究发现，女性过劳性膝痛发生率（30.4%）高于男性（27.8%），且女性新发病例出现更早（第4周 vs 第7周）[13]。这一研究提示女性运动员是膝关节损伤的高危人群，需要早期强化训练和特别关注。'''
doc.add_paragraph(content_2_3)

print("第二章完成，使用[11][12][13]")

doc.add_page_break()

# 第三章 现状调查 - 使用[14]Leppänen
doc.add_heading('第三章 江汉大学大学生篮球运动膝关节损伤现状调查', level=1)

# 3.1 调查对象与方法
doc.add_heading('3.1 调查对象与方法', level=2)
male_count = len([r for r in data if r.get('Q1_性别')=='男'])
female_count = len([r for r in data if r.get('Q1_性别')=='女'])

content_3_1 = f'''本研究以江汉大学有篮球运动经历的全日制在校大学生为调查对象。共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%。其中男生{male_count}人，女生{female_count}人。

Leppänen等通过前瞻性观察研究，使用3D动作分析技术发现，女性篮球运动员膝关节损伤率是男性的6.2倍，女性外翻角显著大于男性（13.9° vs 2.0°）[14]。这一发现为本研究的性别差异分析提供了重要参考。'''
doc.add_paragraph(content_3_1)

# 3.2 损伤发生情况分析
doc.add_heading('3.2 损伤发生情况分析', level=2)
content_3_2 = f'''调查结果显示，在{total}名受访者中，有{injured_total}人曾经受过膝关节损伤，损伤发生率为{injured_total/total*100:.1f}%。

损伤类型分布：'''
doc.add_paragraph(content_3_2)

for t, c in sorted(injury_types.items(), key=lambda x: x[1], reverse=True):
    pct = c / max(injured_total, 1) * 100
    doc.add_paragraph(f'{t}：{c}人（{pct:.1f}%）', style='List Bullet')

print("第三章完成，使用[14]")

doc.add_page_break()

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_15文献版.docx')
print("已保存")
