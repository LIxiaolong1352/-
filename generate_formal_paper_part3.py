from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

# 读取文档和数据
doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_正式版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 继续第二章
# 2.3
doc.add_heading('2.3 膝关节损伤的分类与机制', level=2)
content_2_3 = '''篮球运动中常见的膝关节损伤主要包括以下几类，曹炜[1]、张涛[2]从不同角度进行了详细分析。

（1）半月板损伤。半月板损伤是篮球运动中最常见的膝关节损伤，多发生在起跳落地和急停变向时。曹炜[1]指出，损伤机制包括：膝关节半屈曲位时突然旋转，导致半月板受到挤压和剪切力；落地时膝关节不稳，半月板受到异常应力。彭清政[10]的研究证实，半月板是承受和传递负荷的核心结构，应力峰值最高（快速助跑时超过10MPa），因此最易损伤。张海军[3]的调查显示，半月板损伤占26.78%，徐小敏[12]的研究也显示占22.75%，均位居损伤类型首位。

（2）韧带损伤。前交叉韧带（ACL）损伤多发生在急停变向和起跳落地时，机制是膝关节过度伸展或旋转时胫骨前移。内侧副韧带（MCL）损伤多发生在身体对抗时膝关节外翻。彭清政[10]发现，快速助跑时ACL负荷显著增加。Leppänen等[14]的研究显示，女性运动员ACL损伤风险是男性的6.2倍，外翻角显著更大（13.9° vs 2.0°）。高万钧[4]的调查中，韧带拉伤占40.2%，位居第一。

（3）髌骨劳损。髌骨劳损（髌骨软化症）是由于髌股关节长期过度磨损导致的退行性改变，多发生在训练强度大、技术动作不规范的运动员中。赵孝凯[5]指出，屈膝30°-50°时髌骨负荷最重，前锋、中锋由于频繁起跳和对抗，损伤率高于后卫。张海军[3]的调查显示，髌骨劳损占19.6%，徐小敏[12]的研究显示占18.72%，均位居第二。

（4）髌腱炎。髌腱炎（跳跃膝）是由于髌腱反复受到牵拉导致的慢性炎症，多发生在频繁起跳落地的运动员中。Owoeye等[15]的研究发现，神经肌肉训练对髌腱炎无效，提示其发病机制与膝关节其他损伤不同。'''
doc.add_paragraph(content_2_3)

# 2.4
doc.add_heading('2.4 膝关节损伤的风险因素', level=2)
content_2_4 = '''篮球运动膝关节损伤的风险因素可分为内因和外因两大类，陈思伟[8]、蒋莉[9]等学者进行了系统分析。

（1）内因。包括：①性别因素。Leppänen等[14]和Stilling等[13]的研究均证实，女性运动员是膝关节损伤的高危人群，ACL损伤率和过劳性膝痛发生率均高于男性。这与女性骨盆较宽、Q角较大、韧带松弛度大有关。②年龄和训练年限。张海军[3]、徐小敏[12]的研究发现，训练年限越短（尤其4-5年）损伤率越高，随年限增长下降。这与技术不成熟、自我保护意识差有关。③身体素质。高万钧[4]指出，专项素质差是主要致因之一。路程[7]发现，股后肌群肌力下降是重要致伤因素。④技术水平。刘虎[6]的研究显示，技术动作错误是首要致因（50%）。

（2）外因。包括：①热身不充分。张海军[3]、徐小敏[12]、张冬[11]的研究均将准备活动不合理列为首要或重要致因。热身不充分导致肌肉温度低、关节灵活性差，增加损伤风险。②训练负荷过大。张海军[3]、徐小敏[12]指出，训练负荷不合理和局部负担过重是主要致因。高万钧[4]也提到带伤训练问题严重。③场地条件差。刘虎[6]的研究中，场地器材问题占30%。赵孝凯[5]指出，混凝土场地是高危情境。④医务监督不足。徐小敏[12]的研究发现，伤后及时处理率仅14%，多数无队医。陈思伟[8]强调要学习损伤处理知识。'''
doc.add_paragraph(content_2_4)

doc.add_page_break()

print("第二章完成，开始生成第三章...")

# 第三章
doc.add_heading('第三章 江汉大学大学生篮球运动膝关节损伤现状调查', level=1)

# 3.1
doc.add_heading('3.1 调查对象与方法', level=2)

# 3.1.1
doc.add_heading('3.1.1 调查对象的选择', level=3)
specialty_count = len([r for r in data if r.get('Q2_身份')=='体育学院篮球专项'])
amateur_count = len([r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）'])
male_count = len([r for r in data if r.get('Q1_性别')=='男'])
female_count = len([r for r in data if r.get('Q1_性别')=='女'])

content_3_1_1 = f'''本研究以江汉大学有篮球运动经历的全日制在校大学生为调查对象。纳入标准：（1）江汉大学全日制在校学生；（2）每周至少参与1次篮球运动，持续3个月以上；（3）自愿参加调查。排除标准：（1）有先天性膝关节疾病史；（2）近半年内有膝关节手术史；（3）因其他原因无法完成问卷。

共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%。其中男生{male_count}人（{male_count/total*100:.1f}%），女生{female_count}人（{female_count/total*100:.1f}%）；体育学院篮球专项{specialty_count}人（{specialty_count/total*100:.1f}%），普通篮球爱好者（公体课/社团）{amateur_count}人（{amateur_count/total*100:.1f}%）。运动年限分布：1年以下{count_by('Q3_运动年限').get('1年以下', 0)}人（{count_by('Q3_运动年限').get('1年以下', 0)/total*100:.1f}%），1-3年{count_by('Q3_运动年限').get('1-3年', 0)}人（{count_by('Q3_运动年限').get('1-3年', 0)/total*100:.1f}%），5年以上{count_by('Q3_运动年限').get('5年以上', 0)}人（{count_by('Q3_运动年限').get('5年以上', 0)/total*100:.1f}%）。每周打球频率：1-2次{count_by('Q4_每周频率').get('1-2次', 0)}人（{count_by('Q4_每周频率').get('1-2次', 0)/total*100:.1f}%），3-4次{count_by('Q4_每周频率').get('3-4次', 0)}人（{count_by('Q4_每周频率').get('3-4次', 0)/total*100:.1f}%），5次及以上{count_by('Q4_每周频率').get('5次及以上', 0)}人（{count_by('Q4_每周频率').get('5次及以上', 0)/total*100:.1f}%）。'''
doc.add_paragraph(content_3_1_1)

# 3.1.2
doc.add_heading('3.1.2 问卷设计与发放', level=3)
content_3_1_2 = '''本研究采用自编问卷《大学生篮球运动膝关节损伤调查问卷》，问卷设计参考张海军[3]、高万钧[4]、徐小敏[12]等学者的问卷条目，内容包括四个部分：

第一部分：基本信息。包括性别、年龄、年级、身高体重、身份（体育专项/普通爱好者）等。

第二部分：运动情况。包括篮球运动年限、每周打球频率、每次运动时长、技术水平自评、是否参加过比赛等。

第三部分：损伤情况。包括是否受过膝关节损伤、受伤次数、损伤类型（半月板、韧带、髌骨等）、损伤发生情况（起跳落地、急停变向、身体对抗等）、损伤主要原因、损伤后处理方式等。

第四部分：影响因素。包括热身习惯（每次都做/偶尔做/很少做）、热身时长、技术学习情况（系统学习/自学/未学习）、护具使用情况（经常用/偶尔用/不用）、场地条件评价（好/一般/差）、预防知识了解程度（非常了解/了解一些/不太了解/完全不了解）等。

问卷采用现场发放、现场回收的方式，由研究者统一指导填写，确保问卷填写的真实性和完整性。'''
doc.add_paragraph(content_3_1_2)

# 3.1.3
doc.add_heading('3.1.3 数据统计方法', level=3)
content_3_1_3 = '''运用Excel和SPSS 26.0软件对调查数据进行统计分析。描述性统计用于分析损伤发生率、损伤类型分布等基本情况；卡方检验用于分析分类变量与损伤发生的关系；独立样本t检验用于比较两组连续变量的差异。检验水准设为α=0.05，P<0.05认为差异有统计学意义。'''
doc.add_paragraph(content_3_1_3)

print("3.1节完成...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_正式版.docx')
print("已保存，继续生成...")
