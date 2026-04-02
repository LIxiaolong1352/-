import subprocess
import sys

# 安装必要的库
subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import csv

# 读取CSV数据
with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
print(f"总样本量: {total}")

# 统计函数
def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts, len(filtered)

# 计算统计数据
gender_stats, _ = count_by('Q1_性别')
injury_stats, _ = count_by('Q5_是否受伤')
injured_count = injury_stats.get('是', 0)
injured_rate = (injured_count / total * 100) if total > 0 else 0

warmup_stats, _ = count_by('Q10_准备活动')
knowledge_stats, _ = count_by('Q12_预防知识')
action_stats, _ = count_by('Q13_受伤后做法')

injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)

# 身份分布
identity_stats, _ = count_by('Q2_身份')
# 运动年限
years_stats, _ = count_by('Q3_运动年限')
# 每周频率
freq_stats, _ = count_by('Q4_每周频率')
# 受伤次数
times_stats, _ = count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤') == '是')
# 损伤情况
situation_stats, _ = count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤') == '是')
# 损伤原因
reason_stats, _ = count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤') == '是')
# 活动时长
duration_stats, _ = count_by('Q11_活动时长')

# 损伤类型统计（多选）
injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 读取原始论文 - 使用通配符查找文件
import glob
docx_files = glob.glob(r'C:\Users\Administrator\Desktop\*10000*.docx')
if not docx_files:
    docx_files = glob.glob(r'C:\Users\Administrator\Desktop\*论文*.docx')
if not docx_files:
    raise FileNotFoundError("找不到论文Word文件")

doc_path = docx_files[0]
print(f"找到论文文件: {doc_path}")
doc = Document(doc_path)

# 替换文本函数
def replace_text_in_doc(doc, old_text, new_text):
    for para in doc.paragraphs:
        if old_text in para.text:
            para.text = para.text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)

# 进行替换
replacements = [
    ("50份问卷，回收48份有效问卷", f"{total}份问卷，回收{total}份有效问卷"),
    ("48份有效问卷", f"{total}份有效问卷"),
    ("48名调查对象", f"{total}名受访者"),
    ("48个参与调查的学生", f"{total}个参与调查的学生"),
    ("48人中", f"{total}人中"),
    ("48人中有28人", f"{total}人中有{injured_count}人"),
    ("受伤率达到58.3%", f"受伤率达到{injured_rate:.1f}%"),
    ("损伤率为58.3%", f"损伤率为{injured_rate:.1f}%"),
    ("58.3%的人", f"{injured_rate:.1f}%的人"),
    ("男生36人（75%）", f"男生{gender_stats.get('男', 0)}人（{gender_stats.get('男', 0)/total*100:.1f}%）"),
    ("女生12人（25%）", f"女生{gender_stats.get('女', 0)}人（{gender_stats.get('女', 0)/total*100:.1f}%）"),
    ("男女比例3:1", f"男女比例{gender_stats.get('男', 0)/max(gender_stats.get('女', 1), 1):.2f}:1"),
]

for old, new in replacements:
    replace_text_in_doc(doc, old, new)
    print(f"替换: {old} -> {new}")

# 保存修改后的文档
output_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_已修改.docx'
doc.save(output_path)
print(f"\n修改完成！已保存到: {output_path}")

# 生成数据摘要
summary = f"""
【数据更新摘要】
总样本量: {total}人
男生: {gender_stats.get('男', 0)}人 ({gender_stats.get('男', 0)/total*100:.1f}%)
女生: {gender_stats.get('女', 0)}人 ({gender_stats.get('女', 0)/total*100:.1f}%)
受伤人数: {injured_count}人
损伤率: {injured_rate:.1f}%

身份分布:
- 体育学院篮球专项: {identity_stats.get('体育学院篮球专项', 0)}人
- 普通篮球爱好者: {identity_stats.get('普通篮球爱好者（公体课/社团）', 0)}人

损伤类型:
"""
for t, c in sorted(injury_types.items(), key=lambda x: x[1], reverse=True):
    summary += f"- {t}: {c}人 ({c/max(injured_total,1)*100:.1f}%)\n"

print(summary)
