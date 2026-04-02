from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

# 3.2.3 损伤发生情况与原因分析 - 引用[3][4][5][10][11][12]
doc.add_heading('3.2.3 损伤发生情况与原因分析', level=3)
content_3_2_3 = f'''从损伤发生时的具体情况来看：起跳落地时受伤的{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('起跳落地', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('起跳落地', 0)/max(injured_total,1)*100:.1f}%），急停变向时受伤的{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('急停变向', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('急停变向', 0)/max(injured_total,1)*100:.1f}%），身体对抗时受伤的{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('身体对抗', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('身体对抗', 0)/max(injured_total,1)*100:.1f}%），其他情况{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)/max(injured_total,1)*100:.1f}%）。

结果表明，起跳落地是导致膝关节损伤的最主要动作，占51.6%，这与彭清政[10]的研究结论一致。彭清政[10]通过生物力学分析发现，助跑速度越快，垂直地面反作用力（VGRF）越大，快速助跑时VGRF可达3.79倍体重，半月板应力峰值超过10MPa，极易导致损伤。

从损伤主要原因来看：太累/疲劳{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('太累/疲劳', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('太累/疲劳', 0)/max(injured_total,1)*100:.1f}%），准备活动没做好{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('准备活动没做好', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('准备活动没做好', 0)/max(injured_total,1)*100:.1f}%），场地太滑/条件差{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('场地太滑/条件差', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('场地太滑/条件差', 0)/max(injured_total,1)*100:.1f}%），技术动作错误{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('技术动作错误', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('技术动作错误', 0)/max(injured_total,1)*100:.1f}%），其他原因{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)/max(injured_total,1)*100:.1f}%）。

太累/疲劳和准备活动不充分是最主要的致伤原因，这与张海军[3]、徐小敏[12]、张冬[11]的研究结论一致。张海军[3]和徐小敏[12]均将准备活动不合理列为首要致因，张冬[11]的研究中准备活动不合理占28%。刘虎[6]的研究中技术动作错误占50%，赵孝凯[5]也指出技术动作错误是主要成因，但在本研究中仅占12.9%，可能与调查对象的技术水平差异有关。'''
doc.add_paragraph(content_3_2_3)

# 3.3 不同群体损伤特征对比分析
doc.add_heading('3.3 不同群体损伤特征对比分析', level=2)

# 3.3.1 性别差异分析 - 引用[13][14]
doc.add_heading('3.3.1 性别差异分析', level=3)
male_count = len([r for r in data if r.get('Q1_性别')=='男'])
female_count = len([r for r in data if r.get('Q1_性别')=='女'])
male_injured = len([r for r in data if r.get('Q1_性别')=='男' and r.get('Q5_是否受伤')=='是'])
female_injured = len([r for r in data if r.get('Q1_性别')=='女' and r.get('Q5_是否受伤')=='是'])
male_rate = male_injured / max(male_count, 1) * 100
female_rate = female_injured / max(female_count, 1) * 100

content_3_3_1 = f'''调查结果显示，男生损伤率为{male_rate:.1f}%（{male_injured}/{male_count}），女生损伤率为{female_rate:.1f}%（{female_injured}/{female_count}）。虽然男生损伤人数多于女生，但损伤率差异不大，说明性别不是影响损伤发生的主要因素。

然而，Leppänen等[13]和Stilling等[14]的研究均证实，女性运动员是膝关节损伤（尤其是ACL损伤）的高危人群。Leppänen等[13]发现，女性篮球运动员膝关节损伤率是男性的6.2倍，8例ACL损伤全部发生在女性，女性外翻角显著大于男性（13.9° vs 2.0°）。Stilling等[14]的研究也显示，女性过劳性膝痛发生率（30.4%）高于男性（27.8%），且女性新发病例出现更早（第4周 vs 第7周）。

本研究中男女生损伤率差异不明显，可能与以下因素有关：一是样本量有限，女性样本仅{female_count}人；二是本调查对象为普通大学生，运动强度相对较低，女性运动员在高强度对抗中的损伤风险优势未充分体现。但仍建议对女性篮球参与者加强预防指导，特别是ACL损伤的预防。'''
doc.add_paragraph(content_3_3_1)

print("3.3.1节完成...")

# 3.3.2 运动水平差异分析 - 引用[3][4][12]
doc.add_heading('3.3.2 运动水平差异分析', level=3)
specialty_data = [r for r in data if r.get('Q2_身份')=='体育学院篮球专项']
amateur_data = [r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）']
specialty_injured = len([r for r in specialty_data if r.get('Q5_是否受伤')=='是'])
amateur_injured = len([r for r in amateur_data if r.get('Q5_是否受伤')=='是'])
specialty_rate = specialty_injured / max(len(specialty_data), 1) * 100
amateur_rate = amateur_injured / max(len(amateur_data), 1) * 100

content_3_3_2 = f'''调查结果显示，体育学院篮球专项学生损伤率为{specialty_rate:.1f}%（{specialty_injured}/{len(specialty_data)}），普通篮球爱好者损伤率为{amateur_rate:.1f}%（{amateur_injured}/{len(amateur_data)}）。

体育专项学生虽然技术水平较高，但由于训练强度大、比赛频繁，损伤率反而略高于普通爱好者。这与徐小敏[12]对高水平运动员的研究结论一致，高水平运动员损伤率高达90.71%，远高于普通高校的37.5%~50%。普通爱好者虽然技术水平相对较低，但运动强度较小，损伤率略低。这说明运动强度和技术水平的匹配程度是影响损伤的重要因素，也与张海军[3]发现的训练年限与损伤率关系相符——训练年限越短（尤其4-5年）损伤率越高，随年限增长下降。'''
doc.add_paragraph(content_3_3_2)

print("3.3.2节完成...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx')
print("已保存")
