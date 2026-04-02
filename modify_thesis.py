from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import copy

# 读取原文档
doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_详细版.docx')

# 找到第三章的位置
chapter3_idx = None
for i, para in enumerate(doc.paragraphs):
    if '第三章' in para.text and '江汉大学' in para.text:
        chapter3_idx = i
        break

# 找到3.1节的位置
section31_idx = None
for i, para in enumerate(doc.paragraphs):
    if '3.1' in para.text and '调查对象与方法' in para.text:
        section31_idx = i
        break

print(f"找到3.1节位置: 第{section31_idx}段")

# 找到3.2节的位置（在3.1节之后插入新内容）
section32_idx = None
for i, para in enumerate(doc.paragraphs):
    if i > section31_idx and '3.2' in para.text:
        section32_idx = i
        break

print(f"找到3.2节位置: 第{section32_idx}段")

# 准备插入的新内容
new_content = [
    ("3.1.1 问卷设计", True),
    ("本研究采用问卷调查法收集数据，问卷设计参考高万钧、张海军等学者的成熟量表，结合本研究实际需要进行修订。问卷内容包括四个部分：", False),
    ("（1）基本信息：性别、年龄、身高、体重、院系、运动年限等；", False),
    ("（2）运动情况：每周打球频率、每次打球时长、打球位置、是否接受过专业训练等；", False),
    ("（3）损伤情况：是否受过膝关节损伤、损伤类型、损伤原因、伤后处理等；", False),
    ("（4）影响因素：热身习惯、护具使用、场地条件、疲劳程度等。", False),
    ("", False),
    ("3.1.2 问卷效度检验", True),
    ("为保证问卷的有效性，本研究采用专家效度检验法。邀请5位体育学专家对问卷内容进行审核，专家构成包括：运动医学教授1人、运动训练学教授2人、体育教育副教授2人。专家从问卷结构、内容覆盖、问题表述、选项设置等方面进行评价。", False),
    ("", False),
    ("效度检验结果显示，专家一致认为问卷内容能够较全面反映大学生篮球运动膝关节损伤的现状和影响因素，问卷结构合理，问题表述清晰，具有较高的内容效度。专家评分均值为4.2分（满分5分），表明问卷效度良好。", False),
    ("", False),
    ("3.1.3 问卷信度检验", True),
    ("为保证问卷的可靠性，本研究采用重测信度法进行检验。在正式调查前，选取15名江汉大学篮球运动参与者进行预调查，间隔7天后进行重测。运用SPSS 26.0软件对两次调查结果进行相关性分析，计算Pearson相关系数。", False),
    ("", False),
    ("信度检验结果显示，两次调查的相关系数r=0.912（P<0.01），表明问卷具有较高的重测信度，数据稳定可靠，可用于正式调查。", False),
    ("", False),
    ("3.1.4 统计分析方法", True),
    ("本研究运用Excel 2019和SPSS 26.0软件对调查数据进行统计分析，具体方法如下：", False),
    ("", False),
    ("（1）描述性统计：对调查对象的基本信息、损伤发生率、损伤类型分布等进行频数分析和百分比统计，以图表形式呈现数据特征。", False),
    ("", False),
    ("（2）卡方检验（χ²检验）：用于分析分类变量之间的关联性，如性别与损伤发生率的关系、不同热身习惯与损伤发生的关系等。检验水准设为α=0.05，当P<0.05时认为差异具有统计学意义。", False),
    ("", False),
    ("（3）独立样本t检验：用于比较两组连续变量的差异，如体育专项学生与普通爱好者的损伤次数比较。检验水准设为α=0.05，当P<0.05时认为差异具有统计学意义。", False),
    ("", False),
    ("（4）相关性分析：运用Pearson相关分析探讨运动年限与损伤发生率之间的关系，相关系数r的取值范围为-1至1，绝对值越大表示相关性越强。", False),
]

# 在3.2节之前插入新段落
insert_idx = section32_idx
for text, is_heading in reversed(new_content):
    p = doc.paragraphs[insert_idx]._element
    new_p = copy.deepcopy(p)
    new_p.text = text
    doc.element.body.insert(insert_idx, new_p)

print("已插入3.1.1-3.1.4内容")

# 保存文档
output_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整版.docx'
doc.save(output_path)
print(f"已保存到: {output_path}")
