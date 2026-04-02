import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import re

doc_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整版.docx'
doc = Document(doc_path)

# 清除所有现有引用
for para in doc.paragraphs:
    text = para.text
    new_text = re.sub(r'\[\d+\]', '', text)
    if new_text != text:
        para.text = new_text

print('已清除所有引用')

# 按顺序在指定段落添加引用
# 格式: (段落索引, 引用编号)
citations_in_order = [
    (6, '[1]'),    # 第6段 - 研究背景
    (31, '[2]'),   # 第31段 - 生物力学与解剖学
    (41, '[3]'),   # 第41段 - 损伤率
    (41, '[4]'),   # 第41段 - 普通高校
    (59, '[5]'),   # 第59段 - 损伤原因
    (65, '[6]'),   # 第65段 - 技术动作错误
    (51, '[7]'),   # 第51段 - 运动年限
    (68, '[8]'),   # 第68段 - 预防策略
    (37, '[9]'),   # 第37段 - 高校篮球
    (65, '[10]'),  # 第65段 - 急停起跳
    (61, '[11]'),  # 第61段 - 准备活动
    (45, '[12]'),  # 第45段 - 高水平
    (53, '[13]'),  # 第53段 - 女性
    (47, '[14]'),  # 第47段 - 损伤率
    (76, '[15]'),  # 第76段 - 热身
]

added = []
for para_idx, citation in citations_in_order:
    if para_idx < len(doc.paragraphs):
        para = doc.paragraphs[para_idx]
        para.add_run(citation)
        added.append(citation)
        print(f'已添加 {citation} 到第{para_idx}段')

print(f'\n共添加 {len(added)} 个引用')

# 保存
output_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx'
doc.save(output_path)
print(f'保存完成: {output_path}')

# 验证
doc2 = Document(output_path)
all_citations = []
for para in doc2.paragraphs:
    matches = re.findall(r'\[(\d+)\]', para.text)
    all_citations.extend(matches)

print(f'\n验证: 共 {len(all_citations)} 个引用')
print('引用顺序:', all_citations)

# 检查是否按1-15顺序
expected = [str(i) for i in range(1, 16)]
if all_citations == expected:
    print('✓ 引用正确: 按顺序[1]-[15]')
else:
    print('✗ 顺序不对')
    print('期望:', expected)
    print('实际:', all_citations)
