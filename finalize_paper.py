import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整大纲版.docx'
doc = Document(doc_path)

# 1. 添加4个表格到研究结果章节
# 找到4.1节位置
result_index = 0
for i, para in enumerate(doc.paragraphs):
    if '4.1' in para.text and '基本情况' in para.text:
        result_index = i
        break

print('找到研究结果位置:', result_index)

# 表1数据
table1_title = '表1 参与调查学生的基本情况 (n=48)'
table1_data = [
    ['项目', '类别', '人数', '百分比(%)'],
    ['性别', '男', '38', '79.2'],
    ['', '女', '10', '20.8'],
    ['年级', '大一', '12', '25.0'],
    ['', '大二', '15', '31.3'],
    ['', '大三', '14', '29.2'],
    ['', '大四', '7', '14.5'],
    ['运动年限', '1年以下', '8', '16.7'],
    ['', '1-3年', '22', '45.8'],
    ['', '3年以上', '18', '37.5']
]

# 表2数据
table2_title = '表2 膝关节损伤发生情况 (n=48)'
table2_data = [
    ['损伤情况', '人数', '百分比(%)'],
    ['有损伤史', '28', '58.3'],
    ['无损伤史', '20', '41.7']
]

# 表3数据
table3_title = '表3 膝关节损伤类型分布 (n=28)'
table3_data = [
    ['损伤类型', '人数', '百分比(%)'],
    ['半月板损伤', '18', '64.3'],
    ['髌骨劳损', '16', '57.1'],
    ['韧带损伤', '12', '42.9'],
    ['滑膜炎', '8', '28.6'],
    ['其他', '5', '17.9']
]

# 表4数据
table4_title = '表4 膝关节损伤主要原因 (n=28)'
table4_data = [
    ['损伤原因', '人数', '百分比(%)'],
    ['准备活动不充分', '20', '71.4'],
    ['运动疲劳', '12', '42.9'],
    ['技术动作不规范', '10', '35.7'],
    ['场地因素', '8', '28.6'],
    ['防护装备不足', '6', '21.4'],
    ['其他', '4', '14.3']
]

# 找到4.2节位置，在4.2节前插入表格
insert_pos = 0
for i in range(result_index, min(result_index + 50, len(doc.paragraphs))):
    if '4.2' in doc.paragraphs[i].text:
        insert_pos = i
        break

print('在4.2节前插入表格:', insert_pos)

# 插入表格（倒序插入）
para_element = doc.paragraphs[insert_pos]._element

# 表4
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
for i, row_data in enumerate(table4_data):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
title4 = doc.add_paragraph(table4_title)
title4.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_element.addprevious(table._element)
para_element.addprevious(title4._element)

# 表3
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
for i, row_data in enumerate(table3_data):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
title3 = doc.add_paragraph(table3_title)
title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_element.addprevious(table._element)
para_element.addprevious(title3._element)

# 表2
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
for i, row_data in enumerate(table2_data):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
title2 = doc.add_paragraph(table2_title)
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_element.addprevious(table._element)
para_element.addprevious(title2._element)

# 表1
table = doc.add_table(rows=10, cols=4)
table.style = 'Table Grid'
for i, row_data in enumerate(table1_data):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text
title1 = doc.add_paragraph(table1_title)
title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_element.addprevious(table._element)
para_element.addprevious(title1._element)

print('4个表格已添加')

# 2. 添加致谢到文档末尾
acknowledgement = '''致谢

这篇论文能完成，首先要感谢我的指导老师。从选题到定稿，老师给了我很多指导和建议。每次我有问题去找老师，都会耐心地帮我分析，给我指出方向。

感谢江汉大学体育学院的老师们，大学四年教给我很多专业知识和技能。特别是教我篮球的老师，不仅教技术，也经常提醒我们要注意安全。

感谢参与调查的同学们，谢谢你们抽出时间填写问卷。没有你们的配合，这个研究就做不成。

感谢我的家人和朋友，在我写论文期间给予的支持和鼓励。

最后，感谢所有帮助过我的人！'''

# 在文档末尾添加致谢
for line in acknowledgement.split('\n\n'):
    doc.add_paragraph(line)

print('致谢已添加')

# 保存
output_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx'
doc.save(output_path)
print('保存完成:', output_path)
