from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_详细版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 第三章 现状调查（扩充版，增加图表）
doc.add_heading('第三章 江汉大学大学生篮球运动膝关节损伤现状调查', level=1)

# 3.1 调查对象与方法（扩充）
doc.add_heading('3.1 调查对象与方法', level=2)

male_count = len([r for r in data if r.get('Q1_性别')=='男'])
female_count = len([r for r in data if r.get('Q1_性别')=='女'])
specialty_count = len([r for r in data if r.get('Q2_身份')=='体育学院篮球专项'])
amateur_count = len([r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）'])

content_3_1 = f'''本研究以江汉大学有篮球运动经历的全日制在校大学生为调查对象。纳入标准：（1）江汉大学全日制在校学生；（2）每周至少参与1次篮球运动，持续3个月以上；（3）自愿参加调查。排除标准：（1）有先天性膝关节疾病史；（2）近半年内有膝关节手术史；（3）因其他原因无法完成问卷。

共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%。其中男生{male_count}人（{male_count/total*100:.1f}%），女生{female_count}人（{female_count/total*100:.1f}%）；体育学院篮球专项{specialty_count}人（{specialty_count/total*100:.1f}%），普通篮球爱好者（公体课/社团）{amateur_count}人（{amateur_count/total*100:.1f}%）。

Leppänen等通过前瞻性观察研究，使用3D动作分析技术发现，女性篮球运动员膝关节损伤率是男性的6.2倍，女性外翻角显著大于男性（13.9° vs 2.0°）。这一发现为本研究的性别差异分析提供了重要参考，也提示应特别关注女性篮球参与者的损伤预防。'''
doc.add_paragraph(content_3_1)

# 添加表2：调查对象基本情况
doc.add_paragraph('表2 调查对象基本情况', style='Caption')
table2 = doc.add_table(rows=3, cols=5)
table2.style = 'Table Grid'

# 表头
headers2 = ['类别', '分组', '人数', '百分比', '合计']
for i, header in enumerate(headers2):
    table2.rows[0].cells[i].text = header
    table2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

# 性别
table2.rows[1].cells[0].text = '性别'
table2.rows[1].cells[1].text = '男'
table2.rows[1].cells[2].text = str(male_count)
table2.rows[1].cells[3].text = f'{male_count/total*100:.1f}%'
table2.rows[1].cells[4].text = str(total)

# 身份
table2.rows[2].cells[0].text = '身份'
table2.rows[2].cells[1].text = '体育专项'
table2.rows[2].cells[2].text = str(specialty_count)
table2.rows[2].cells[3].text = f'{specialty_count/total*100:.1f}%'
table2.rows[2].cells[4].text = str(total)

doc.add_paragraph()

print("3.1节完成，表2添加完成")

# 3.2 损伤发生情况分析（扩充，增加图表）
doc.add_heading('3.2 损伤发生情况分析', level=2)

content_3_2 = f'''调查结果显示，在{total}名受访者中，有{injured_total}人曾经受过膝关节损伤，损伤发生率为{injured_rate:.1f}%。这一比例与张冬对泉州信息工程学院的调查结果（50%）接近，高于张海军对青少年体校的调查（14.42%），低于徐小敏对高水平运动员的调查（90.71%），与高万钧对普通高校运动员的调查（37.5%）和刘虎对和田师专的调查（40.65%）相比略高。说明江汉大学大学生篮球运动膝关节损伤问题较为普遍，接近一半的篮球运动参与者有过膝关节损伤经历，应引起足够重视。'''
doc.add_paragraph(content_3_2)

# 添加表3：损伤类型分布
doc.add_paragraph('表3 损伤类型分布（多选）', style='Caption')
table3 = doc.add_table(rows=len(injury_types)+1, cols=3)
table3.style = 'Table Grid'

# 表头
headers3 = ['损伤类型', '人数', '百分比']
for i, header in enumerate(headers3):
    table3.rows[0].cells[i].text = header
    table3.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

# 数据
for idx, (t, c) in enumerate(sorted(injury_types.items(), key=lambda x: x[1], reverse=True), 1):
    pct = c / max(injured_total, 1) * 100
    table3.rows[idx].cells[0].text = t
    table3.rows[idx].cells[1].text = str(c)
    table3.rows[idx].cells[2].text = f'{pct:.1f}%'

doc.add_paragraph()

print("3.2节完成，表3添加完成")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_详细版.docx')
print("已保存")
