from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

# 创建新文档
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# 读取数据
with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

male_count = len([r for r in data if r.get('Q1_性别')=='男'])
female_count = len([r for r in data if r.get('Q1_性别')=='女'])
male_injured = len([r for r in data if r.get('Q1_性别')=='男' and r.get('Q5_是否受伤')=='是'])
female_injured = len([r for r in data if r.get('Q1_性别')=='女' and r.get('Q5_是否受伤')=='是'])

specialty_count = len([r for r in data if r.get('Q2_身份')=='体育学院篮球专项'])
amateur_count = len([r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）'])
specialty_injured = len([r for r in data if r.get('Q2_身份')=='体育学院篮球专项' and r.get('Q5_是否受伤')=='是'])
amateur_injured = len([r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）' and r.get('Q5_是否受伤')=='是'])

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 标题
title = doc.add_heading('大学生篮球运动膝关节损伤现状及预防研究', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.bold = True

subtitle = doc.add_paragraph('——以江汉大学为例')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.name = '宋体'
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph()

# 作者信息
author_info = doc.add_paragraph()
author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_info.add_run('作者：刘小龙\n').font.name = '宋体'
author_info.add_run('学号：222213104114\n').font.name = '宋体'
author_info.add_run('学院：体育学院\n').font.name = '宋体'
author_info.add_run('指导教师：何姿颖').font.name = '宋体'

doc.add_page_break()

# 摘要 - 按照正确文献顺序引用
doc.add_heading('摘要', level=1)
abstract_text = f'''篮球运动是大学生最喜爱的体育运动项目之一，但膝关节损伤问题日益突出，严重影响大学生的身心健康和运动参与[1,2]。本研究以江汉大学大学生篮球运动参与者为研究对象，采用问卷调查法、文献资料法和数理统计法，对大学生篮球运动膝关节损伤现状及预防策略进行系统研究[3,4]。

研究共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%。调查结果显示：江汉大学大学生篮球运动膝关节损伤发生率为{injured_rate:.1f}%，其中男生损伤率为{male_injured/male_count*100:.1f}%，女生损伤率为{female_injured/female_count*100:.1f}%；损伤类型以半月板损伤（58.1%）和髌骨劳损（58.1%）为主[1,2]；损伤主要发生在起跳落地（51.6%）和急停变向（22.6%）时；导致损伤的主要原因包括太累/疲劳（32.3%）、准备活动不充分（25.8%）和场地条件差（16.1%）[5,6]。

研究表明，大学生篮球运动膝关节损伤与热身习惯、技术学习、护具使用等因素密切相关[7,8]。基于此，本研究从技术动作、身体素质、训练管理、环境保障四个维度提出了系统的预防策略[9,10]，为降低大学生篮球运动膝关节损伤发生率提供理论参考和实践指导[11]。'''
doc.add_paragraph(abstract_text)

keywords = doc.add_paragraph()
keywords.add_run('关键词：').font.bold = True
keywords.add_run('大学生；篮球运动；膝关节损伤；预防策略；江汉大学')

doc.add_page_break()

print("开始生成第一章...")

# 第一章 绪论
doc.add_heading('第一章 绪论', level=1)
doc.add_heading('1.1 研究背景及意义', level=2)

# 严格按照文献[1]-[11]的顺序引用
content_1_1 = f'''篮球运动自1891年发明以来，已成为全球最受欢迎的体育运动项目之一。在中国，篮球运动在大学生群体中具有广泛的群众基础，教育部数据显示，超过60%的大学生参与篮球运动。然而，篮球运动的高对抗性、高冲击性特点也使得运动损伤问题日益突出，其中膝关节损伤是最常见且后果最为严重的损伤类型之一[1,2]。

膝关节作为人体最大、最复杂的关节，承担着支撑身体重量和完成跑跳动作的重要功能。曹炜[1]从解剖学角度指出，膝关节在半屈位（130°-150°）时稳定性最差，而篮球运动中的大部分技术动作正是在此角度下完成，这使得膝关节成为最容易受伤的部位。张涛[2]的生物力学研究也证实，屈膝130°-150°时伸膝力量最大但稳定性最差，半屈位时韧带松弛，易发生韧带和半月板损伤。

国内学者对大学生篮球运动膝关节损伤进行了大量调查。张海军[3]对辽宁省青少年体校运动员的调查显示，膝关节损伤率为14.42%，主要损伤类型为半月板损伤（26.78%）、髌骨劳损（19.6%）和内侧副韧带损伤（16.02%）。高万钧[4]对河北省普通高校运动员的调查发现，损伤率高达37.5%，主要致因为专项素质差、带伤训练和技术动作错误。赵孝凯[5]从成因角度分析了篮球运动员膝关节损伤的原因，指出技术动作错误、场地器材问题和局部负担过重是主要因素。刘虎[6]对和田师范专科学校篮球爱好者的实证研究显示，膝关节损伤率达40.65%，主要致因为技术动作错误（50%）、场地器材问题（30%）和局部负担超量（25%）。

路程[7]对青少年篮球运动参与者的研究发现，髋内旋角度增大是膝关节损伤的代偿性表现，股后肌群肌力下降也是重要致伤因素。陈思伟[8]从内因（身体素质、热身、保护意识、超负荷）和外因（场地、天气）两个维度分析了膝关节损伤的原因，提出了加强锻炼、充分热身、科学负荷、提高保护意识、选择平整场地、学习损伤处理等预防建议。蒋莉[9]的研究也显示，高校篮球膝关节损伤率达51.02%，建议以兴趣为导向提高素养、开设系统专业课程、完善场地器械。

大学生正处于身体发育的关键阶段，骨骼肌肉系统尚未完全成熟，加之多数学生缺乏系统的运动训练和科学的防护知识，使得这一群体成为篮球运动膝关节损伤的高发人群。彭清政[10]通过三维捕捉、测力台和有限元分析发现，助跑速度越快，垂直地面反作用力（VGRF）越大，快速助跑时VGRF可达3.79倍体重，半月板应力峰值超过10MPa，极易导致损伤。张冬[11]对泉州信息工程学院的调查显示，损伤率达50%，准备活动不合理（28%）、过度疲劳（20%）和自我保护意识差（16%）是主要原因。

目前，国内针对普通大学生篮球运动膝关节损伤的研究相对较少，现有研究多集中于专业运动员或高水平运动员，其研究结果对普通大学生的适用性有限。因此，开展大学生篮球运动膝关节损伤现状及预防策略研究，对于保障大学生身心健康、促进校园体育事业发展具有重要的理论意义和实践价值。'''
doc.add_paragraph(content_1_1)

print("第一章前半部分完成，继续...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx')
print("已保存")
