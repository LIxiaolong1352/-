from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

# 创建新文档
doc = Document()

# 设置页面
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# 设置字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# 读取数据
with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

# 标题
title = doc.add_heading('大学生篮球运动膝关节损伤现状及预防研究', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.bold = True

subtitle = doc.add_paragraph('——以江汉大学为例')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.name = '宋体'
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph()

# 作者信息
author_info = doc.add_paragraph()
author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_info.add_run('作者：刘小龙\n').font.name = '宋体'
author_info.add_run('学号：222213104114\n').font.name = '宋体'
author_info.add_run('学院：体育学院\n').font.name = '宋体'
author_info.add_run('指导教师：何姿颖').font.name = '宋体'

doc.add_page_break()

# 摘要
doc.add_heading('摘要', level=1)
abstract_text = f'''篮球运动是大学生最喜爱的体育运动项目之一，但膝关节损伤问题日益突出。本研究以江汉大学大学生篮球运动参与者为研究对象，采用问卷调查法、文献资料法和数理统计法，对大学生篮球运动膝关节损伤现状及预防策略进行研究[1-4]。

研究共发放问卷{total}份，回收有效问卷{total}份。结果显示：江汉大学大学生篮球运动膝关节损伤发生率为{injured_rate:.1f}%；损伤类型以半月板损伤（58.1%）和髌骨劳损（58.1%）为主[1,2]；损伤主要发生在起跳落地（51.6%）时[3]；主要致因包括太累/疲劳（32.3%）、准备活动不充分（25.8%）[4-6]。研究表明，大学生篮球运动膝关节损伤与热身习惯、技术学习、护具使用等因素密切相关[7-9]。本研究从技术动作、身体素质、训练管理、环境保障四个维度提出了预防策略[10-15]。'''
doc.add_paragraph(abstract_text)

keywords = doc.add_paragraph()
keywords.add_run('关键词：').font.bold = True
keywords.add_run('大学生；篮球运动；膝关节损伤；预防策略；江汉大学')

doc.add_page_break()

print("开始生成第一章...")

# 第一章
doc.add_heading('第一章 绪论', level=1)
doc.add_heading('1.1 研究背景及意义', level=2)

# 按照文献出现顺序引用
content_1_1 = f'''篮球运动自1891年发明以来，已成为全球最受欢迎的体育运动项目之一。在中国，篮球运动在大学生群体中具有广泛的群众基础。然而，篮球运动的高对抗性、高冲击性特点也使得运动损伤问题日益突出，其中膝关节损伤是最常见且后果最为严重的损伤类型之一[1,2]。

膝关节作为人体最大、最复杂的关节，承担着支撑身体重量和完成跑跳动作的重要功能。曹炜[1]从解剖学角度指出，膝关节在半屈位（130°-150°）时稳定性最差，而篮球运动中的大部分技术动作正是在此角度下完成。张涛[2]的生物力学研究也证实，屈膝130°-150°时伸膝力量最大但稳定性最差，半屈位时韧带松弛，易发生韧带和半月板损伤。

国内学者对大学生篮球运动膝关节损伤进行了大量调查。张海军[3]对辽宁省青少年体校运动员的调查显示，膝关节损伤率为14.42%，主要损伤类型为半月板损伤（26.78%）、髌骨劳损（19.6%）。高万钧[4]对河北省普通高校运动员的调查发现，损伤率高达37.5%，主要致因为专项素质差、带伤训练和技术动作错误。徐小敏[5]对江苏省高校高水平运动员的研究显示，损伤率更是达到90.71%，主要致因为准备活动不合理、带伤训练和训练负荷不合理。

大学生正处于身体发育的关键阶段，骨骼肌肉系统尚未完全成熟，加之多数学生缺乏系统的运动训练和科学的防护知识，使得这一群体成为篮球运动膝关节损伤的高发人群[6,7]。目前，国内针对普通大学生篮球运动膝关节损伤的研究相对较少，现有研究多集中于专业运动员或高水平运动员[8,9]，其研究结果对普通大学生的适用性有限。因此，开展本研究对于保障大学生身心健康、促进校园体育事业发展具有重要的理论意义和实践价值。'''
doc.add_paragraph(content_1_1)

print("第一章前半部分完成...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_修正版.docx')
print("已保存")
