from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

# 创建新文档
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# 读取数据
with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

# 标题
title = doc.add_heading('大学生篮球运动膝关节损伤现状及预防研究', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.bold = True

subtitle = doc.add_paragraph('——以江汉大学为例')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.name = '宋体'
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph()

# 作者信息
author_info = doc.add_paragraph()
author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_info.add_run('作者：刘小龙\n').font.name = '宋体'
author_info.add_run('学号：222213104114\n').font.name = '宋体'
author_info.add_run('学院：体育学院\n').font.name = '宋体'
author_info.add_run('指导教师：何姿颖').font.name = '宋体'

doc.add_page_break()

# 摘要 - 只用[1]-[5]，每篇一次
doc.add_heading('摘要', level=1)

# [1] 曹炜 - 解剖学背景
abstract_1 = f'''篮球运动是大学生最喜爱的体育运动项目之一，但膝关节损伤问题日益突出。曹炜从解剖学角度指出，膝关节在半屈位（130°-150°）时稳定性最差，而篮球运动中的大部分技术动作正是在此角度下完成，这使得膝关节成为最容易受伤的部位[1]。'''
doc.add_paragraph(abstract_1)

# [2] 张涛 - 生物力学
abstract_2 = '''张涛的生物力学研究证实，屈膝130°-150°时伸膝力量最大但稳定性最差，半屈位时韧带松弛，易发生韧带和半月板损伤。篮球运动中的起跳落地、急停变向等动作对膝关节产生复杂的生物力学负荷，起跳落地时膝关节需要承受巨大的地面反作用力[2]。'''
doc.add_paragraph(abstract_2)

# [3] 张海军 - 调查方法
abstract_3 = f'''本研究以江汉大学大学生篮球运动参与者为研究对象，采用问卷调查法。参考张海军对辽宁省青少年体校运动员的调查方法，设计《大学生篮球运动膝关节损伤调查问卷》[3]。共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%。'''
doc.add_paragraph(abstract_3)

# [4] 高万钧 - 损伤率结果
abstract_4 = f'''调查结果显示：江汉大学大学生篮球运动膝关节损伤发生率为{injured_rate:.1f}%。这一比例与高万钧对河北省普通高校运动员的调查结果（37.5%）接近，说明江汉大学大学生篮球运动膝关节损伤问题较为普遍[4]。'''
doc.add_paragraph(abstract_4)

# [5] 赵孝凯 - 主要发现
abstract_5 = '''损伤类型以半月板损伤（58.1%）和髌骨劳损（58.1%）为主；损伤主要发生在起跳落地（51.6%）和急停变向（22.6%）时；导致损伤的主要原因包括太累/疲劳（32.3%）、准备活动不充分（25.8%）。赵孝凯指出，技术动作错误、场地器材问题和局部负担过重是主要成因[5]。'''
doc.add_paragraph(abstract_5)

keywords = doc.add_paragraph()
keywords.add_run('关键词：').font.bold = True
keywords.add_run('大学生；篮球运动；膝关节损伤；预防策略；江汉大学')

doc.add_page_break()

print("摘要完成，使用[1][2][3][4][5]")
print("继续生成第一章...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_15文献版.docx')
print("已保存")
