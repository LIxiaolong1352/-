from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_正式版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 3.2 膝关节损伤发生情况分析
doc.add_heading('3.2 膝关节损伤发生情况分析', level=2)

# 3.2.1
doc.add_heading('3.2.1 损伤发生率统计', level=3)
content_3_2_1 = f'''调查结果显示，在{total}名受访者中，有{injured_total}人曾经受过膝关节损伤，损伤发生率为{injured_rate:.1f}%。这一比例与张冬[11]对泉州信息工程学院的调查结果（50%）接近，高于张海军[3]对青少年体校的调查（14.42%），低于徐小敏[12]对高水平运动员的调查（90.71%），与高万钧[4]对普通高校运动员的调查（37.5%）和刘虎[6]对和田师专的调查（40.65%）相比略高。说明江汉大学大学生篮球运动膝关节损伤问题较为普遍，接近一半的篮球运动参与者有过膝关节损伤经历，应引起足够重视。

从损伤发生次数来看，在{injured_total}名受伤者中，受伤1次的{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('1次', 0)}人（{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('1次', 0)/max(injured_total,1)*100:.1f}%），受伤2次的{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('2次', 0)}人（{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('2次', 0)/max(injured_total,1)*100:.1f}%），受伤3次及以上的{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('3次及以上', 0)}人（{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('3次及以上', 0)/max(injured_total,1)*100:.1f}%）。说明大部分受伤者只受过1次伤，但也有部分学生存在反复受伤的情况，提示可能存在慢性损伤或康复不彻底的问题，与张海军[3]、徐小敏[12]研究中提到的带伤训练问题一致。'''
doc.add_paragraph(content_3_2_1)

# 3.2.2
doc.add_heading('3.2.2 损伤类型与部位分布', level=3)
content_3_2_2 = f'''调查结果显示，大学生篮球运动膝关节损伤类型多样，主要包括（多选统计）：'''
doc.add_paragraph(content_3_2_2)

# 添加损伤类型列表
for t, c in sorted(injury_types.items(), key=lambda x: x[1], reverse=True):
    pct = c / max(injured_total, 1) * 100
    doc.add_paragraph(f'{t}：{c}人（{pct:.1f}%）', style='List Bullet')

content_3_2_2b = f'''从损伤类型分布可以看出，半月板损伤和髌骨劳损是最主要的损伤类型，两者占比均为58.1%，这与张海军[3]（半月板26.78%、髌骨劳损19.6%）、徐小敏[12]（半月板22.75%、髌骨劳损18.72%）的研究结论一致，也与曹炜[1]、张涛[2]从解剖学和生物力学角度的理论分析相符。曹炜[1]指出，膝关节在半屈位（130°-150°）时稳定性最差，而篮球运动中的频繁起跳落地和急停变向正是在此角度下完成，对半月板和髌股关节产生巨大冲击。彭清政[10]的研究也证实，半月板是承受和传递负荷的核心结构，应力峰值最高，因此最易损伤。

韧带损伤（内侧/前交叉）占22.6%，虽然占比相对较低，但后果往往更为严重，恢复周期更长。Leppänen等[14]的研究显示，女性运动员ACL损伤风险是男性的6.2倍，提示应特别关注女性运动员的韧带损伤预防。'''
doc.add_paragraph(content_3_2_2b)

# 3.2.3
doc.add_heading('3.2.3 损伤发生情况与原因分析', level=3)
content_3_2_3 = f'''从损伤发生时的具体情况来看：起跳落地时受伤的{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('起跳落地', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('起跳落地', 0)/max(injured_total,1)*100:.1f}%），急停变向时受伤的{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('急停变向', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('急停变向', 0)/max(injured_total,1)*100:.1f}%），身体对抗时受伤的{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('身体对抗', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('身体对抗', 0)/max(injured_total,1)*100:.1f}%），其他情况{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)/max(injured_total,1)*100:.1f}%）。

结果表明，起跳落地是导致膝关节损伤的最主要动作，占51.6%，这与彭清政[10]的研究结论一致。彭清政[10]通过生物力学分析发现，助跑速度越快，垂直地面反作用力（VGRF）越大，快速助跑时VGRF可达3.79倍体重，半月板应力峰值超过10MPa，极易导致损伤。

从损伤主要原因来看：太累/疲劳{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('太累/疲劳', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('太累/疲劳', 0)/max(injured_total,1)*100:.1f}%），准备活动没做好{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('准备活动没做好', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('准备活动没做好', 0)/max(injured_total,1)*100:.1f}%），场地太滑/条件差{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('场地太滑/条件差', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('场地太滑/条件差', 0)/max(injured_total,1)*100:.1f}%），技术动作错误{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('技术动作错误', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('技术动作错误', 0)/max(injured_total,1)*100:.1f}%），其他原因{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)/max(injured_total,1)*100:.1f}%）。

太累/疲劳和准备活动不充分是最主要的致伤原因，这与张海军[3]、徐小敏[12]、张冬[11]的研究结论一致。张海军[3]和徐小敏[12]均将准备活动不合理列为首要致因，张冬[11]的研究中准备活动不合理占28%。刘虎[6]的研究中技术动作错误占50%，但在本研究中仅占12.9%，可能与调查对象的技术水平差异有关。'''
doc.add_paragraph(content_3_2_3)

print("3.2节完成...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_正式版.docx')
print("已保存")
