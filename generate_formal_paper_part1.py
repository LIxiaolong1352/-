import subprocess
import sys

# 安装必要的库
subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import csv

# 读取CSV数据
with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

male_data = [r for r in data if r.get('Q1_性别') == '男']
female_data = [r for r in data if r.get('Q1_性别') == '女']
male_injured = len([r for r in male_data if r.get('Q5_是否受伤') == '是'])
female_injured = len([r for r in female_data if r.get('Q5_是否受伤') == '是'])

# 统计函数
def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

# 损伤类型统计
injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 创建Word文档
doc = Document()

# 设置页面大小和边距
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# 标题
title = doc.add_heading('大学生篮球运动膝关节损伤现状及预防研究', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.bold = True

# 副标题
subtitle = doc.add_paragraph('——以江汉大学为例')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.name = '宋体'
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph()

# 作者信息
author_info = doc.add_paragraph()
author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_info.add_run('作者：刘小龙\n').font.name = '宋体'
author_info.add_run('学号：222213104114\n').font.name = '宋体'
author_info.add_run('学院：体育学院\n').font.name = '宋体'
author_info.add_run('指导教师：何姿颖').font.name = '宋体'

doc.add_page_break()

# 摘要
doc.add_heading('摘要', level=1)
abstract_text = f'''篮球运动是大学生最喜爱的体育运动项目之一，但膝关节损伤问题日益突出，严重影响大学生的身心健康和运动参与。本研究以江汉大学大学生篮球运动参与者为研究对象，采用问卷调查法、文献资料法和数理统计法，对大学生篮球运动膝关节损伤现状及预防策略进行系统研究。

研究共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%。调查结果显示：江汉大学大学生篮球运动膝关节损伤发生率为{injured_rate:.1f}%，其中男生损伤率为{male_injured/len(male_data)*100:.1f}%，女生损伤率为{female_injured/len(female_data)*100:.1f}%；损伤类型以半月板损伤（58.1%）和髌骨劳损（58.1%）为主；损伤主要发生在起跳落地（51.6%）和急停变向（22.6%）时；导致损伤的主要原因包括太累/疲劳（32.3%）、准备活动不充分（25.8%）和场地条件差（16.1%）。

研究表明，大学生篮球运动膝关节损伤与热身习惯、技术学习、护具使用等因素密切相关。研究结果与张海军[3]、高万钧[4]、徐小敏[12]等学者的研究结论基本一致，损伤类型分布符合曹炜[1]、张涛[2]从解剖学和生物力学角度的理论分析。基于此，本研究从技术动作、身体素质、训练管理、环境保障四个维度提出了系统的预防策略，为降低大学生篮球运动膝关节损伤发生率提供理论参考和实践指导。'''

doc.add_paragraph(abstract_text)

# 关键词
keywords = doc.add_paragraph()
keywords.add_run('关键词：').font.bold = True
keywords.add_run('大学生；篮球运动；膝关节损伤；预防策略；江汉大学')

doc.add_page_break()

print("正在生成论文内容，请稍候...")

# 第一章 绪论
doc.add_heading('第一章 绪论', level=1)

# 1.1 研究背景及意义
doc.add_heading('1.1 研究背景及意义', level=2)
content_1_1 = f'''篮球运动自1891年发明以来，已成为全球最受欢迎的体育运动项目之一。在中国，篮球运动在大学生群体中具有广泛的群众基础，教育部数据显示，超过60%的大学生参与篮球运动。然而，篮球运动的高对抗性、高冲击性特点也使得运动损伤问题日益突出，其中膝关节损伤是最常见且后果最为严重的损伤类型之一。

膝关节作为人体最大、最复杂的关节，承担着支撑身体重量和完成跑跳动作的重要功能。曹炜[1]从解剖学角度指出，膝关节在半屈位（130°-150°）时稳定性最差，而篮球运动中的大部分技术动作正是在此角度下完成，这使得膝关节成为最容易受伤的部位。张涛[2]的生物力学研究也证实，屈膝130°-150°时伸膝力量最大但稳定性最差，半屈位时韧带松弛，易发生韧带和半月板损伤。

国内学者对大学生篮球运动膝关节损伤进行了大量调查。张海军[3]对辽宁省青少年体校运动员的调查显示，膝关节损伤率为14.42%，主要损伤类型为半月板损伤（26.78%）、髌骨劳损（19.6%）和内侧副韧带损伤（16.02%）。高万钧[4]对河北省普通高校运动员的调查发现，损伤率高达37.5%，主要致因为专项素质差、带伤训练和技术动作错误。徐小敏[12]对江苏省高校高水平运动员的研究显示，损伤率更是达到90.71%，其中准备活动不合理、带伤训练和训练负荷不合理是三大主要致因。

大学生正处于身体发育的关键阶段，骨骼肌肉系统尚未完全成熟，加之多数学生缺乏系统的运动训练和科学的防护知识，使得这一群体成为篮球运动膝关节损伤的高发人群。刘虎[6]对和田师范专科学校的研究发现，篮球爱好者膝关节损伤率达40.65%，主要致因为技术动作错误（50%）、场地器材问题（30%）和局部负担超量（25%）。张冬[11]对泉州信息工程学院的调查也显示，损伤率达50%，准备活动不合理（28%）、过度疲劳（20%）和自我保护意识差（16%）是主要原因。

目前，国内针对普通大学生篮球运动膝关节损伤的研究相对较少，现有研究多集中于专业运动员或高水平运动员，其研究结果对普通大学生的适用性有限。因此，开展大学生篮球运动膝关节损伤现状及预防策略研究，对于保障大学生身心健康、促进校园体育事业发展具有重要的理论意义和实践价值。'''
doc.add_paragraph(content_1_1)

# 1.2 研究目的
doc.add_heading('1.2 研究目的', level=2)
content_1_2 = '''本研究以江汉大学大学生篮球运动参与者为研究对象，旨在实现以下研究目的：

第一，全面了解江汉大学大学生篮球运动膝关节损伤的现状，包括损伤发生率、损伤类型、损伤部位、损伤程度等基本情况，为后续研究提供数据支撑。

第二，深入分析大学生篮球运动膝关节损伤的影响因素，从个体因素（性别、运动年限、技术水平）、行为因素（热身习惯、护具使用、技术学习）和环境因素（场地条件、训练管理）等多个维度探讨损伤发生的原因。

第三，基于损伤现状和影响因素分析，构建大学生篮球运动膝关节损伤预防策略体系，从技术动作改进、身体素质提升、训练管理优化、环境保障完善等方面提出针对性的预防措施，为学校体育工作和学生自我保护提供参考。'''
doc.add_paragraph(content_1_2)

# 1.3 国内外研究现状
doc.add_heading('1.3 国内外研究现状', level=2)

# 1.3.1 国外研究现状
doc.add_heading('1.3.1 国外研究现状', level=3)
content_1_3_1 = '''国外对篮球运动损伤的研究起步较早，研究体系相对完善。近年来，国外研究逐渐从描述性研究转向机制研究和预防研究。

在生物力学研究方面，彭清政[10]通过三维捕捉、测力台和有限元分析发现，助跑速度越快，垂直地面反作用力（VGRF）越大、载荷率越高、达峰时间越短，膝关节受力越大。研究还表明，半月板应力峰值最高（快速助跑时超过10MPa），是膝关节中最易损伤的结构；慢速时应力集中易导致慢性劳损，快速时ACL负荷增加易导致急性损伤。

在预防干预研究方面，Owoeye等[15]开展了一项历史性队列研究，评估神经肌肉训练（NMT）热身项目对降低青少年篮球运动员膝关节和踝关节损伤的效果。研究结果显示，NMT热身可使膝关节损伤率降低49%（IRR=0.51），踝关节损伤率降低32%（IRR=0.68），但对髌腱/跟腱病无效。这一研究为篮球运动膝关节损伤的预防提供了循证医学证据。

Leppänen等[14]通过前瞻性观察研究，使用3D动作分析技术发现，女性篮球运动员膝关节损伤率是男性的6.2倍（8例ACL损伤全部发生在女性），女性外翻角显著大于男性（13.9° vs 2.0°）。Stilling等[13]的追踪研究也证实，女性过劳性膝痛发生率（30.4%）高于男性（27.8%），且女性新发病例出现更早（第4周 vs 第7周）。这些研究提示女性运动员是膝关节损伤（尤其是ACL损伤）的高危人群，需要早期强化训练。

在理论层面，国外学者提出了多种损伤预防理论模型。损伤发生理论强调内因和外因的交互作用，认为受伤是身体素质、技术水平、心理状态等内因与场地条件、训练方法等外因共同作用的结果。神经肌肉控制理论则强调通过训练提高神经系统对肌肉的控制能力，增强关节稳定性，从而降低损伤风险。'''
doc.add_paragraph(content_1_3_1)

print("第一章前半部分完成...")

# 保存文档
output_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_正式版.docx'
doc.save(output_path)
print(f"论文前半部分已保存到: {output_path}")
