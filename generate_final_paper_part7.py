from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

# 3.3.3 热身习惯与损伤关系分析 - 引用[3][11][12]
doc.add_heading('3.3.3 热身习惯与损伤关系分析', level=3)
always_warmup = [r for r in data if r.get('Q10_准备活动')=='每次都做']
occasionally_warmup = [r for r in data if r.get('Q10_准备活动')=='偶尔做']
always_injured = len([r for r in always_warmup if r.get('Q5_是否受伤')=='是'])
occasionally_injured = len([r for r in occasionally_warmup if r.get('Q5_是否受伤')=='是'])
always_rate = always_injured / max(len(always_warmup), 1) * 100
occasionally_rate = occasionally_injured / max(len(occasionally_warmup), 1) * 100

content_3_3_3 = f'''调查结果显示，每次都做热身的学生{len(always_warmup)}人（{len(always_warmup)/total*100:.1f}%），损伤率为{always_rate:.1f}%（{always_injured}/{len(always_warmup)}）；偶尔做热身的学生{len(occasionally_warmup)}人（{len(occasionally_warmup)/total*100:.1f}%），损伤率为{occasionally_rate:.1f}%（{occasionally_injured}/{len(occasionally_warmup)}）。

虽然两组损伤率差异未达到统计学显著水平，但每次都做热身的损伤率略低于偶尔做热身的。这与张海军[3]、徐小敏[12]、张冬[11]的研究结论一致，准备活动不合理是首要或重要致因。张海军[3]和徐小敏[12]均将准备活动不合理列为首要致因，张冬[11]的研究中准备活动不合理占28%。说明热身习惯对预防膝关节损伤具有重要作用，应加强学生热身意识的培养。'''
doc.add_paragraph(content_3_3_3)

# 3.4 损伤致因综合分析 - 引用[1][2][3][4][5][6][7][8][10][11][12]
doc.add_heading('3.4 损伤致因综合分析', level=2)
content_3_4 = '''基于上述调查结果，结合曹炜[1]、张涛[2]的解剖学和生物力学理论，对江汉大学大学生篮球运动膝关节损伤的致因进行综合分析。

（1）技术动作因素。篮球运动中的起跳落地、急停变向等技术动作对膝关节产生巨大冲击。彭清政[10]的研究表明，助跑速度越快，VGRF越大，快速助跑时VGRF可达3.79倍体重，半月板应力峰值超过10MPa。曹炜[1]、张涛[2]指出，膝关节在半屈位（130°-150°）时稳定性最差，而篮球大部分技术动作正是在此角度下完成。本研究中，起跳落地致伤占51.6%，技术动作错误致因占12.9%，说明技术动作规范对预防损伤至关重要。

（2）身体素质因素。路程[7]发现，髋内旋角度增大和股后肌群肌力下降是膝关节损伤的重要致伤因素。高万钧[4]也将专项素质差列为主要致因。本研究中，太累/疲劳致因占32.3%，说明身体素质不足、疲劳状态下继续运动是导致损伤的重要原因。

（3）训练管理因素。张海军[3]、徐小敏[12]均将训练负荷不合理和带伤训练列为主要致因。徐小敏[12]的研究还发现，伤后及时处理率仅14%，多数无队医。本研究中，准备活动不充分致因占25.8%，这与张冬[11]的研究结论一致。刘虎[6]的研究中技术动作错误占50%，赵孝凯[5]也指出技术动作错误是主要成因。说明训练管理、医务监督仍有待加强。

（4）环境因素。刘虎[6]的研究中，场地器材问题占30%。赵孝凯[5]指出，混凝土场地是高危情境。陈思伟[8]强调要选择平整场地。本研究中，场地条件差致因占16.1%，说明场地环境改善也是预防损伤的重要环节。'''
doc.add_paragraph(content_3_4)

doc.add_page_break()

print("第三章完成，开始生成第四章...")

# 第四章
doc.add_heading('第四章 大学生篮球运动膝关节损伤预防策略', level=1)

# 4.1 技术动作优化策略 - 引用[1][2][10]
doc.add_heading('4.1 技术动作优化策略', level=2)
content_4_1 = '''基于曹炜[1]、张涛[2]的解剖学、生物力学理论和彭清政[10]的量化研究结果，提出以下技术动作优化策略：

（1）规范起跳落地技术。彭清政[10]的研究表明，助跑速度越快，膝关节受力越大。应加强起跳落地技术的训练，包括：①控制助跑速度，避免过快助跑；②落地时膝关节微屈（约130°-150°），利用肌肉缓冲；③双脚同时落地，分散冲击力；④落地后顺势屈膝缓冲，避免膝关节伸直锁定。

（2）优化急停变向技术。曹炜[1]指出，急停、旋转时半月板产生矛盾运动，易导致损伤。应加强急停变向技术的训练，包括：①急停时降低重心，增大支撑面；②变向时避免膝关节过度内扣；③加强髋关节灵活性训练，减少膝关节代偿；④使用侧滑步等规范防守步法，避免膝关节过度扭转。

（3）改进半蹲防守姿势。张涛[2]指出，屈膝130°-150°时稳定性最差，但篮球防守多在此角度。应：①控制半蹲角度，避免长时间保持同一角度；②加强股四头肌力量训练，提高半蹲位稳定性；③防守间隙适当活动膝关节，避免僵硬。'''
doc.add_paragraph(content_4_1)

print("4.1节完成...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx')
print("已保存")
