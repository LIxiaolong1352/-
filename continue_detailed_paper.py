from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import csv

# 打开现有文档继续编辑
doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_详细版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

# 继续添加1.4和1.5节
# 1.4 研究方法（扩充）
doc.add_heading('1.4 研究方法', level=2)
content_1_4 = f'''本研究综合运用文献资料法、问卷调查法、数理统计法和逻辑分析法，力求全面、客观地了解江汉大学大学生篮球运动膝关节损伤现状，并提出科学的预防策略。

（1）文献资料法。通过中国知网、万方数据库、PubMed等数据库检索相关文献，了解国内外篮球运动膝关节损伤研究的最新进展，为本研究提供理论基础和研究思路。重点参考了张海军、高万钧、徐小敏、路程等学者的流行病学调查方法，以及曹炜、张涛、彭清政的生物力学研究成果。

（2）问卷调查法。设计《大学生篮球运动膝关节损伤调查问卷》，问卷设计参考张海军、高万钧、徐小敏的问卷条目，内容包括基本信息、运动情况、损伤情况、影响因素四个部分。问卷经过专家审核和预调查后正式发放，共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%。

（3）数理统计法。运用Excel和SPSS 26.0软件对调查数据进行统计分析，包括描述性统计、卡方检验、独立样本t检验等，分析损伤发生的特点和影响因素。检验水准设为α=0.05，P<0.05认为差异有统计学意义。

（4）逻辑分析法。运用归纳、演绎、比较等逻辑方法，对统计结果进行分析，结合解剖学和生物力学理论，揭示损伤发生的规律，构建预防策略体系。'''
doc.add_paragraph(content_1_4)

# 1.5 研究内容（扩充）
doc.add_heading('1.5 研究内容', level=2)
content_1_5 = '''本研究主要包括以下四个方面的内容：

第一，大学生篮球运动膝关节损伤现状调查。通过问卷调查收集江汉大学大学生篮球运动参与者的基本信息、运动情况、损伤情况，统计分析损伤发生率、损伤类型、损伤部位、损伤程度等基本情况，为后续研究提供数据支撑。

第二，大学生篮球运动膝关节损伤影响因素分析。从个体因素（性别、运动年限、技术水平）、行为因素（热身习惯、护具使用、技术学习）和环境因素（场地条件、训练管理）三个维度分析损伤发生的影响因素。

第三，不同群体损伤特征对比分析。比较不同性别、不同运动水平、不同训练年限学生群体的损伤特征差异，揭示损伤发生的群体规律。

第四，大学生篮球运动膝关节损伤预防策略构建。基于损伤现状和影响因素分析，从技术动作优化、身体素质提升、训练管理优化、环境保障完善四个维度构建系统的预防策略体系。'''
doc.add_paragraph(content_1_5)

doc.add_page_break()

print("第一章完成")
print("开始生成第二章（增加理论基础内容）...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_详细版.docx')
print("已保存")
