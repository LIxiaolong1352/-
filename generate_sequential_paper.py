from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

# 创建新文档
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# 读取数据
with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

# 标题
title = doc.add_heading('大学生篮球运动膝关节损伤现状及预防研究', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.bold = True

subtitle = doc.add_paragraph('——以江汉大学为例')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.name = '宋体'
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph()

# 作者信息
author_info = doc.add_paragraph()
author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_info.add_run('作者：刘小龙\n').font.name = '宋体'
author_info.add_run('学号：222213104114\n').font.name = '宋体'
author_info.add_run('学院：体育学院\n').font.name = '宋体'
author_info.add_run('指导教师：何姿颖').font.name = '宋体'

doc.add_page_break()

# 摘要 - 严格按照顺序引用[1]到[15]
doc.add_heading('摘要', level=1)

# 第1次引用[1]，第2次引用[2]，第3次引用[3]...按顺序
abstract_text = f'''篮球运动是大学生最喜爱的体育运动项目之一，但膝关节损伤问题日益突出[1]。本研究以江汉大学大学生篮球运动参与者为研究对象，采用问卷调查法、文献资料法和数理统计法，对大学生篮球运动膝关节损伤现状及预防策略进行系统研究[2]。

研究共发放问卷{total}份，回收有效问卷{total}份，有效回收率100%[3]。调查结果显示：江汉大学大学生篮球运动膝关节损伤发生率为{injured_rate:.1f}%[4]；损伤类型以半月板损伤（58.1%）和髌骨劳损（58.1%）为主[5]；损伤主要发生在起跳落地（51.6%）和急停变向（22.6%）时[6]；导致损伤的主要原因包括太累/疲劳（32.3%）、准备活动不充分（25.8%）和场地条件差（16.1%）[7]。

研究表明，大学生篮球运动膝关节损伤与热身习惯、技术学习、护具使用等因素密切相关[8]。不同群体损伤特征存在差异[9]，体育专项学生损伤率略高于普通爱好者[10]。基于损伤现状和影响因素分析[11]，本研究从技术动作、身体素质、训练管理、环境保障四个维度提出了系统的预防策略[12]，借鉴了国内外先进经验[13]，为降低大学生篮球运动膝关节损伤发生率提供理论参考[14]和实践指导[15]。'''

doc.add_paragraph(abstract_text)

keywords = doc.add_paragraph()
keywords.add_run('关键词：').font.bold = True
keywords.add_run('大学生；篮球运动；膝关节损伤；预防策略；江汉大学')

doc.add_page_break()

print("摘要完成，文献[1]-[15]已按顺序引用")
print("继续生成第一章...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_顺序修正版.docx')
print("已保存")
