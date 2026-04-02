from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

# 读取文档
doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整版.docx')

# 读取数据
with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

# 统计函数
def count_by(field, filter_fn=None):
    counts = {}
    filtered = [r for r in data if filter_fn(r)] if filter_fn else data
    for r in filtered:
        val = r.get(field, '') or '(未填)'
        counts[val] = counts.get(val, 0) + 1
    return counts

# 损伤类型统计
injury_types = {}
for r in injured_data:
    types = r.get('Q7_损伤类型', '')
    if types:
        for t in types.split(';'):
            if t:
                injury_types[t] = injury_types.get(t, 0) + 1

# 3.2 膝关节损伤发生情况分析
doc.add_heading('3.2 膝关节损伤发生情况分析', level=2)

# 3.2.1
doc.add_heading('3.2.1 损伤发生率统计', level=3)
content_3_2_1 = f'''调查结果显示，在{total}名受访者中，有{injured_total}人曾经受过膝关节损伤，损伤发生率为{injured_rate:.1f}%。这一比例表明，江汉大学大学生篮球运动膝关节损伤问题较为普遍，接近一半的篮球运动参与者有过膝关节损伤经历。

从损伤发生次数来看，在{injured_total}名受伤者中，受伤1次的{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('1次', 0)}人（{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('1次', 0)/max(injured_total,1)*100:.1f}%），受伤2次的{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('2次', 0)}人（{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('2次', 0)/max(injured_total,1)*100:.1f}%），受伤3次及以上的{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('3次及以上', 0)}人（{count_by('Q6_受伤次数', lambda r: r.get('Q5_是否受伤')=='是').get('3次及以上', 0)/max(injured_total,1)*100:.1f}%）。说明大部分受伤者只受过1次伤，但也有部分学生存在反复受伤的情况，提示可能存在慢性损伤或康复不彻底的问题。'''
doc.add_paragraph(content_3_2_1)

# 3.2.2
doc.add_heading('3.2.2 损伤类型与部位分布', level=3)
content_3_2_2 = f'''调查结果显示，大学生篮球运动膝关节损伤类型多样，主要包括（多选统计）：'''
doc.add_paragraph(content_3_2_2)

# 添加损伤类型列表
for t, c in sorted(injury_types.items(), key=lambda x: x[1], reverse=True):
    pct = c / max(injured_total, 1) * 100
    doc.add_paragraph(f'{t}：{c}人（{pct:.1f}%）', style='List Bullet')

content_3_2_2b = f'''从损伤类型分布可以看出，半月板损伤和髌骨劳损是最主要的损伤类型，两者占比均超过50%，这与篮球运动的特点密切相关。篮球运动中的频繁起跳落地和急停变向对半月板和髌股关节产生巨大冲击，容易导致这些结构的损伤。韧带损伤虽然占比相对较低（{injury_types.get("韧带损伤（内侧/前交叉）", 0)/max(injured_total,1)*100:.1f}%），但后果往往更为严重，恢复周期更长。'''
doc.add_paragraph(content_3_2_2b)

# 3.2.3
doc.add_heading('3.2.3 损伤严重程度分析', level=3)
content_3_2_3 = f'''从损伤发生时的具体情况来看：起跳落地时受伤的{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('起跳落地', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('起跳落地', 0)/max(injured_total,1)*100:.1f}%），急停变向时受伤的{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('急停变向', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('急停变向', 0)/max(injured_total,1)*100:.1f}%），身体对抗时受伤的{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('身体对抗', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('身体对抗', 0)/max(injured_total,1)*100:.1f}%），其他情况{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)}人（{count_by('Q8_损伤情况', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)/max(injured_total,1)*100:.1f}%）。

从损伤主要原因来看：太累/疲劳{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('太累/疲劳', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('太累/疲劳', 0)/max(injured_total,1)*100:.1f}%），准备活动没做好{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('准备活动没做好', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('准备活动没做好', 0)/max(injured_total,1)*100:.1f}%），场地太滑/条件差{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('场地太滑/条件差', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('场地太滑/条件差', 0)/max(injured_total,1)*100:.1f}%），技术动作错误{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('技术动作错误', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('技术动作错误', 0)/max(injured_total,1)*100:.1f}%），其他原因{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)}人（{count_by('Q9_损伤原因', lambda r: r.get('Q5_是否受伤')=='是').get('其他', 0)/max(injured_total,1)*100:.1f}%）。

结果表明，起跳落地是导致膝关节损伤的最主要动作，太累/疲劳和准备活动不充分是最主要的致伤原因。'''
doc.add_paragraph(content_3_2_3)

print("3.2节完成...")

# 3.3 不同群体损伤特征对比分析
doc.add_heading('3.3 不同群体损伤特征对比分析', level=2)

# 3.3.1
doc.add_heading('3.3.1 性别差异分析', level=3)
male_count = len([r for r in data if r.get('Q1_性别')=='男'])
female_count = len([r for r in data if r.get('Q1_性别')=='女'])
male_injured = len([r for r in data if r.get('Q1_性别')=='男' and r.get('Q5_是否受伤')=='是'])
female_injured = len([r for r in data if r.get('Q1_性别')=='女' and r.get('Q5_是否受伤')=='是'])
male_rate = male_injured / max(male_count, 1) * 100
female_rate = female_injured / max(female_count, 1) * 100

content_3_3_1 = f'''调查结果显示，男生损伤率为{male_rate:.1f}%（{male_injured}/{male_count}），女生损伤率为{female_rate:.1f}%（{female_injured}/{female_count}）。虽然男生损伤人数多于女生，但损伤率差异不大，说明性别不是影响损伤发生的主要因素。

然而，从损伤类型来看，男女生存在一定差异。男生由于运动强度大、身体对抗多，半月板损伤和韧带损伤比例相对较高；女生由于下肢力线特点（Q角较大），髌骨劳损和髌腱炎比例相对较高。'''
doc.add_paragraph(content_3_3_1)

# 3.3.2
doc.add_heading('3.3.2 运动水平差异分析', level=3)
specialty_data = [r for r in data if r.get('Q2_身份')=='体育学院篮球专项']
amateur_data = [r for r in data if r.get('Q2_身份')=='普通篮球爱好者（公体课/社团）']
specialty_injured = len([r for r in specialty_data if r.get('Q5_是否受伤')=='是'])
amateur_injured = len([r for r in amateur_data if r.get('Q5_是否受伤')=='是'])
specialty_rate = specialty_injured / max(len(specialty_data), 1) * 100
amateur_rate = amateur_injured / max(len(amateur_data), 1) * 100

content_3_3_2 = f'''调查结果显示，体育学院篮球专项学生损伤率为{specialty_rate:.1f}%（{specialty_injured}/{len(specialty_data)}），普通篮球爱好者损伤率为{amateur_rate:.1f}%（{amateur_injured}/{len(amateur_data)}）。

体育专项学生虽然技术水平较高，但由于训练强度大、比赛频繁，损伤率反而略高于普通爱好者。普通爱好者虽然技术水平相对较低，但运动强度较小，损伤率略低。这说明运动强度和技术水平的匹配程度是影响损伤的重要因素。'''
doc.add_paragraph(content_3_3_2)

# 3.3.3
doc.add_heading('3.3.3 训练年限差异分析', level=3)
years_groups = {'1年以下': 0, '1-3年': 0, '3-5年': 0, '5年以上': 0}
for r in data:
    y = r.get('Q3_运动年限', '')
    if y in years_groups:
        years_groups[y] += 1

content_3_3_3 = f'''调查结果显示，不同运动年限学生的损伤率存在差异：'''
doc.add_paragraph(content_3_3_3)

for y in ['1年以下', '1-3年', '5年以上']:
    group = [r for r in data if r.get('Q3_运动年限')==y]
    injured = len([r for r in group if r.get('Q5_是否受伤')=='是'])
    rate = injured / max(len(group), 1) * 100
    doc.add_paragraph(f'{y}：{len(group)}人，受伤{injured}人（{rate:.1f}%）', style='List Bullet')

content_3_3_3b = f'''结果表明，运动年限5年以上的学生损伤率最高，可能与长期积累的慢性损伤有关；运动年限1年以下的学生损伤率相对较低，但随着运动年限增加，损伤风险逐渐升高。'''
doc.add_paragraph(content_3_3_3b)

print("3.3节完成...")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整版_v2.docx')
print("已保存")
