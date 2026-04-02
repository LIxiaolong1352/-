from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# 读取 Markdown 文件
with open('C:\\Users\\Administrator\\.openclaw\\workspace\\毕业论文_刘小龙_10000字版.md', 'r', encoding='utf-8') as f:
    content = f.read()

# 创建 Word 文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

# 添加标题
title = doc.add_heading('大学生篮球运动膝关节损伤现状及预防研究', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.name = '黑体'
title_run.font.size = Pt(18)
title_run.font.bold = True

# 添加副标题
subtitle = doc.add_paragraph('——以江苏大学为例')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.name = '宋体'
subtitle_run.font.size = Pt(14)

# 添加作者信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_text = '''作者：李希龙
学号：222213104114
学院：教师教育学院
指导教师：何姿颖
完成日期：2026年3月15日'''
info.add_run(info_text).font.name = '宋体'
info.runs[0].font.size = Pt(12)

# 添加分页符
doc.add_page_break()

# 处理正文内容
lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    
    # 跳过已处理的标题部分
    if i < 15:
        i += 1
        continue
    
    # 处理标题
    if line.startswith('# '):
        heading = doc.add_heading(line[2:], level=1)
        heading.runs[0].font.name = '黑体'
        heading.runs[0].font.size = Pt(16)
    elif line.startswith('## '):
        heading = doc.add_heading(line[3:], level=2)
        heading.runs[0].font.name = '黑体'
        heading.runs[0].font.size = Pt(14)
    elif line.startswith('### '):
        heading = doc.add_heading(line[4:], level=3)
        heading.runs[0].font.name = '黑体'
        heading.runs[0].font.size = Pt(12)
    elif line.startswith('**') and line.endswith('**'):
        # 粗体文本
        p = doc.add_paragraph()
        run = p.add_run(line[2:-2])
        run.font.bold = True
        run.font.name = '宋体'
    elif line.startswith('- ') or line.startswith('（'):
        # 列表项
        p = doc.add_paragraph(line, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
    elif line.startswith('[') and ']' in line:
        # 参考文献
        p = doc.add_paragraph(line)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
    elif line == '---':
        # 分隔线，添加空行
        doc.add_paragraph()
    elif line:
        # 普通段落
        p = doc.add_paragraph(line)
        p.paragraph_format.first_line_indent = Inches(0.5)
    
    i += 1

# 保存文档
output_path = 'C:\\Users\\Administrator\\Desktop\\毕业论文_刘小龙_10000字.docx'
doc.save(output_path)
print(f'Word 文档已保存到: {output_path}')
