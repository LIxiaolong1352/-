from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx')

# 5.2 研究局限与展望 - 引用[3][7][10][12][13][14]
doc.add_heading('5.2 研究局限与展望', level=2)
content_5_2 = '''本研究存在以下局限：

（1）样本量有限。本研究仅调查了江汉大学68名篮球运动参与者，样本量较小，代表性有限。未来研究可扩大样本量，纳入多所高校，提高研究的外部效度。

（2）研究设计为横断面调查，无法确定因果关系。未来研究可采用前瞻性队列研究设计，追踪观察损伤发生情况，更深入地分析影响因素。

（3）未进行干预效果验证。本研究提出的预防策略基于文献分析和理论推导，未进行实证检验。未来研究可开展干预试验，验证预防策略的实际效果，特别是Owoeye等[12]提出的神经肌肉训练（NMT）在普通大学生中的适用性。

（4）未纳入生物力学指标。未来研究可借鉴彭清政[10]、路程[7]的方法，使用三维动作捕捉、测力台等设备，量化分析大学生篮球运动中的膝关节受力特征，为预防策略提供更精确的依据。

（5）女性样本量不足。本研究女性样本仅24人，未能充分分析性别差异。未来研究应增加女性样本，深入分析女性运动员ACL损伤的高危因素，借鉴Leppänen等[13]、Stilling等[14]的研究方法，制定针对性的预防方案。'''
doc.add_paragraph(content_5_2)

doc.add_page_break()

print("开始生成参考文献...")

# 参考文献 - 严格按照提供的顺序
doc.add_heading('参考文献', level=1)

references = [
    "[1] 曹炜. 从解剖学角度分析篮球运动中膝关节损伤及预防[J]. 湖北科技学院学报, 2013, 33(12): 162-163.",
    "[2] 张涛. 篮球运动中膝关节损伤的生物力学特征分析[C]//中国体育科学学会运动生物力学分会. 第十四届全国运动生物力学学术交流大会论文集. 成都体育学院, 2010: 454-456.",
    "[3] 张海军. 辽宁省青少年篮球运动员膝关节损伤现状调查与对策研究[D]. 辽宁师范大学, 2011.",
    "[4] 高万钧. 河北省普通高校篮球运动员运动损伤流行病学调查及致因研究[D]. 河北师范大学, 2011.",
    "[5] 赵孝凯. 浅谈篮球运动员膝关节损伤成因及措施[J]. 青春岁月, 2012.",
    "[6] 刘虎, 陆玉坤. 篮球运动膝关节损伤的实证研究——以和田师范专科学校为例[J]. 和田师范专科学校学报, 2016, 35(05): 58-61.",
    "[7] 路程. 青少年篮球运动参与者膝关节损伤特点及致伤因素分析[D]. 辽宁师范大学, 2020.",
    "[8] 陈思伟. 论篮球运动膝关节损伤及其预防策略[J]. 冰雪体育创新研究, 2021, (01): 97-98.",
    "[9] 蒋莉. 高校篮球运动中膝关节损伤原因及预防策略的研究[C]//中国体育科学学会运动生物力学分会. 第二十二届全国运动生物力学学术交流大会论文摘要集. 西南医科大学体育学院, 2022: 537-539.",
    "[10] 彭清政. 篮球运动不同速度急停跳中枢脚触地阶段膝关节生物力学特征分析[D]. 武汉体育学院, 2024.",
    "[11] 张冬. 高校篮球运动员膝关节损伤调查与预防研究[J]. 当代体育科技, 2020, 10(26): 31-33.",
    "[12] 徐小敏. 江苏省高校高水平篮球运动员膝关节损伤现状及对策研究[D]. 苏州大学, 2010.",
    "[13] Stilling C, Owoeyea B O, Benson C L, et al. 319 Knee and ankle overuse injuries in youth basketball players[J]. British Journal of Sports Medicine, 2021, 55(Suppl 1): A122-A123.",
    "[14] Leppänen M, Parkkari J, Vasankari T, et al. 042 Change of direction biomechanics and the risk for non-contact knee injuries in youth basketball and floorball players[J]. British Journal of Sports Medicine, 2021, 55(Suppl 1): A17-A18.",
    "[15] Owoeye O, Pasanen K, Befus K, et al. 056 The effectiveness of neuromuscular training warm-up programme to reduce knee and ankle injuries in youth basketball: a historical cohort study[J]. British Journal of Sports Medicine, 2021, 55(Suppl 1): A23-A23."
]

for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.left_indent = Cm(0.75)

print("参考文献完成！")

# 保存最终版本
output_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx'
doc.save(output_path)
print(f"论文全部完成！已保存到: {output_path}")

# 统计字数
import zipfile
from xml.etree import ElementTree

word_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

try:
    docx = zipfile.ZipFile(output_path)
    xml = docx.read('word/document.xml')
    tree = ElementTree.fromstring(xml)
    paragraphs = tree.findall(f'.//{word_ns}p')
    full_text = []
    for p in paragraphs:
        texts = [node.text for node in p.findall(f'.//{word_ns}t') if node.text]
        full_text.append(''.join(texts))
    text = '\n'.join(full_text)
    word_count = len(text)
    print(f"\n论文字数统计：约 {word_count} 字")
except:
    print("\n字数统计失败，但论文已完成")
