import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def markdown_to_docx(md_file, docx_file):
    """将Markdown文件转换为Word文档"""
    
    # 读取Markdown内容
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 按行处理
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # 处理标题
        if line.startswith('# '):
            # 一级标题
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
        elif line.startswith('## '):
            # 二级标题
            doc.add_heading(line[3:], level=2)
            i += 1
        elif line.startswith('### '):
            # 三级标题
            doc.add_heading(line[4:], level=3)
            i += 1
        elif line.startswith('**') and line.endswith('**'):
            # 粗体段落（可能是作者信息等）
            text = line.strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            i += 1
        elif line.startswith('---'):
            # 分隔线，跳过
            i += 1
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                item_text = lines[i].strip()[2:]
                # 处理加粗
                item_text = re.sub(r'\*\*(.*?)\*\*', r'\1', item_text)
                list_items.append(item_text)
                i += 1
            
            for item in list_items:
                p = doc.add_paragraph(item, style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            # 编号列表
            text = re.sub(r'^\d+\.\s', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            i += 1
        elif line.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            # 表格
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) >= 2:
                # 解析表格
                rows = []
                for tl in table_lines:
                    if '---' not in tl:
                        cells = [c.strip() for c in tl.split('|')[1:-1]]
                        if cells:
                            rows.append(cells)
                
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    for row_idx, row_cells in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_cells):
                            table.rows[row_idx].cells[col_idx].text = cell_text
        else:
            # 普通段落
            # 处理加粗
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            
            # 处理引用
            if text.startswith('> '):
                text = text[2:]
            
            doc.add_paragraph(text)
            i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f"转换完成: {md_file} -> {docx_file}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    
    md_file = sys.argv[1]
    docx_file = sys.argv[2]
    markdown_to_docx(md_file, docx_file)
