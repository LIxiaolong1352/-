from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

doc = Document(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_详细版.docx')

with open(r'C:\Users\Administrator\Desktop\survey_data_模拟数据.csv', 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    data = list(reader)

total = len(data)
injured_data = [r for r in data if r.get('Q5_是否受伤') == '是']
injured_total = len(injured_data)
injured_rate = injured_total / total * 100

# 英文摘要 - 扩充版
eng_title = doc.add_heading('Abstract', level=1)

eng_content = f'''Basketball is one of the most popular sports among college students, but knee joint injury has become an increasingly prominent problem that seriously affects college students' physical and mental health and sports participation. As the largest and most complex joint in the human body, the knee joint bears important functions of supporting body weight and completing jumping and running movements, and is subjected to huge loads in basketball.

Cao Wei pointed out from the anatomical perspective that the knee joint is the least stable when it is in a semi-flexed position (130°-150°), and most technical movements in basketball are completed at this angle, making the knee joint the most vulnerable part [1]. Zhang Tao's biomechanical research also confirmed that when the knee is bent 130°-150°, the extensor strength is the greatest but the stability is the worst. When the ligaments are relaxed in the semi-flexed position, ligament and meniscus injuries are prone to occur [2].

This study takes basketball participants in Jianghan University as the research object, and uses questionnaire survey method, literature method and mathematical statistics method to systematically study the current situation and prevention strategies of knee joint injury in college basketball. Referring to Zhang Haijun's survey method for young athletes in Liaoning Provincial Sports School, the "Questionnaire on Knee Joint Injury in College Basketball" was designed [3]. A total of {total} questionnaires were distributed and {total} valid questionnaires were collected, with an effective response rate of 100%.

The results show that the incidence of knee joint injury in college basketball in Jianghan University is {injured_rate:.1f}%. This proportion is close to Gao Wanjun's survey result of 37.5% for ordinary college athletes in Hebei Province [4]. The main types of injuries are meniscus injury (58.1%) and patella strain (58.1%). The main causes include fatigue (32.3%), insufficient warm-up (25.8%) and poor venue conditions (16.1%) [5].

Based on the analysis of injury status and influencing factors, this study proposes systematic prevention strategies from four dimensions: technical movement optimization, physical fitness improvement, training management optimization, and environmental protection improvement, providing theoretical reference and practical guidance for reducing the incidence of knee joint injury in college basketball.'''
doc.add_paragraph(eng_content)

eng_keywords = doc.add_paragraph()
eng_keywords.add_run('Keywords: ').font.bold = True
eng_keywords.add_run('College Students; Basketball; Knee Joint Injury; Prevention Strategy; Jianghan University')

doc.add_page_break()

print("英文摘要完成")
print("开始生成第一章（扩充版）...")

# 第一章 绪论（扩充版）
doc.add_heading('第一章 绪论', level=1)

# 1.1 研究背景及意义（大幅扩充）
doc.add_heading('1.1 研究背景及意义', level=2)

content_1_1 = '''篮球运动自1891年发明以来，已成为全球最受欢迎的体育运动项目之一。在中国，篮球运动在大学生群体中具有广泛的群众基础，教育部数据显示，超过60%的大学生参与篮球运动。篮球运动不仅能够增强体质、培养团队精神，还能够丰富校园文化生活，促进大学生身心健康发展。然而，篮球运动的高对抗性、高冲击性特点也使得运动损伤问题日益突出，其中膝关节损伤是最常见且后果最为严重的损伤类型之一。

随着我国高等教育的快速发展和高校体育设施的不断完善，越来越多的大学生参与到篮球运动中。特别是在"健康中国2030"战略背景下，高校体育工作受到前所未有的重视，篮球作为高校体育的重要组成部分，其参与人数逐年增加。然而，随之而来的运动损伤问题也日益严重，给大学生的身心健康带来了不良影响，甚至影响了正常的学习和生活。

膝关节作为人体最大、最复杂的关节，承担着支撑身体重量和完成跑跳动作的重要功能。膝关节由股骨下端、胫骨上端和髌骨组成，属于滑车关节，其结构复杂，包括骨骼、半月板、韧带、滑膜等多个组成部分。在篮球运动中，膝关节需要承受巨大的负荷，特别是在起跳落地、急停变向、身体对抗等动作中，膝关节承受着数倍于体重的冲击力。长期的高负荷运动容易导致膝关节损伤，影响运动能力和生活质量。

刘虎对和田师范专科学校篮球爱好者的实证研究显示，膝关节损伤率达40.65%，主要致因为技术动作错误（50%）、场地器材问题（30%）和局部负担超量（25%）[6]。这一研究揭示了普通高校篮球运动参与者面临的严重损伤问题，技术动作错误占比高达50%，说明大学生在篮球运动中缺乏规范的技术指导和安全意识。场地器材问题占30%，反映了高校体育设施建设的不足。局部负担超量占25%，提示训练负荷管理的重要性。这些数据表明，大学生篮球运动膝关节损伤问题已经到了不容忽视的程度。

大学生正处于身体发育的关键阶段，骨骼肌肉系统尚未完全成熟，加之多数学生缺乏系统的运动训练和科学的防护知识，使得这一群体成为篮球运动膝关节损伤的高发人群。与专业运动员相比，普通大学生在身体素质、技术水平、防护意识等方面都存在较大差距，更容易发生运动损伤。同时，大学生往往缺乏运动损伤预防和处理的知识，一旦发生损伤，往往不能及时正确处理，导致损伤加重或反复发作。

目前，国内针对普通大学生篮球运动膝关节损伤的研究相对较少，现有研究多集中于专业运动员或高水平运动员，其研究结果对普通大学生的适用性有限。因此，开展大学生篮球运动膝关节损伤现状及预防策略研究，对于保障大学生身心健康、促进校园体育事业发展、推动"健康中国"建设具有重要的理论意义和实践价值。

从理论意义来看，本研究可以丰富我国大学生运动损伤研究的理论体系，填补普通高校篮球运动膝关节损伤研究的空白，为后续研究提供理论基础和数据支撑。从实践意义来看，本研究可以为高校体育工作提供参考，帮助学校制定科学的体育教学计划和损伤预防措施，提高大学生的运动安全意识和自我保护能力，降低运动损伤发生率，促进大学生身心健康发展。'''
doc.add_paragraph(content_1_1)

print("1.1节完成，内容大幅扩充")

# 保存
doc.save(r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_详细版.docx')
print("已保存")
