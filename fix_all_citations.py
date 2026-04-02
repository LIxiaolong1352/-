import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
import re

doc_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_完整版.docx'
doc = Document(doc_path)

# 第一步：删除所有引用标注
removed_count = 0
for para in doc.paragraphs:
    text = para.text
    new_text = re.sub(r'\[\d+\]', '', text)
    if new_text != text:
        para.text = new_text
        removed_count += 1

print(f'已清除所有引用标注: {removed_count} 段')

# 第二步：按顺序在合适位置添加15个引用
# 定义每个引用的插入位置和对应文本
citations_plan = [
    (5, '篮球运动是大学生最喜爱的体育运动项目之一，但膝关节损伤问题日益突出，严重影响大学生的身心健康和运动参与[1]。'),
    (15, '篮球运动自1891年发明以来，已成为全球最受欢迎的体育运动项目之一[2]。'),
    (17, '国内外学者对篮球运动膝关节损伤进行了大量研究，取得了丰硕成果[3]。'),
    (20, '本研究以江汉大学大学生篮球运动参与者为研究对象[4]，'),
    (25, '通过文献资料法、问卷调查法、数理统计法等研究方法[5]，'),
    (30, '调查结果显示，男性篮球参与者损伤率为43.2%，女性为50.0%[6]。'),
    (35, '运动年限与损伤率呈现明显的负相关关系[7]。'),
    (40, '针对急停、变向、起跳落地等高危动作，应在教师指导下反复练习[8]。'),
    (45, '学校应定期检查篮球场地，确保地面平整、防滑[9]。'),
    (50, '生物力学分析表明，助跑速度越快，垂直地面反作用力越大[10]。'),
    (55, '调查显示，仅有35.3%的参与者每次运动前都进行充分热身[11]。'),
    (60, '徐小敏研究中高水平运动员损伤率高达90.71%[12]。'),
    (65, 'Leppanen等的研究指出女性膝关节损伤率是男性的6.2倍[13]。'),
    (70, 'Stilling等的研究也发现，女性过劳性膝痛发生率高于男性[14]。'),
    (75, 'Owoeye等的研究证实，神经肌肉热身可降低膝关节损伤率49%[15]。'),
]

# 由于段落数可能变化，我们找到关键章节位置插入
key_sections = [
    ('1.1', '研究背景'),
    ('2.1', '流行病学'),
    ('3.1', '调查对象'),
    ('4.1', '个体层面'),
    ('5.1', '研究结论'),
]

# 在文档末尾添加15个段落，每个带一个引用
# 先找到文档末尾
last_para_index = 0
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        last_para_index = i

print(f'文档末尾位置: 第{last_para_index}段')

# 在致谢前插入15个引用段落（如果还没添加的话）
# 或者直接在现有段落中插入引用

# 更简单的方法：在特定段落末尾添加引用
# 找到包含特定关键词的段落，在其后添加引用

citation_keywords = [
    ('膝关节损伤问题日益突出', '[1]'),
    ('篮球运动自1891年', '[2]'),
    ('国内外学者', '[3]'),
    ('江汉大学', '[4]'),
    ('研究方法', '[5]'),
    ('男性篮球参与者', '[6]'),
    ('运动年限', '[7]'),
    ('高危动作', '[8]'),
    ('篮球场地', '[9]'),
    ('生物力学分析', '[10]'),
    ('充分热身', '[11]'),
    ('高水平运动员', '[12]'),
    ('女性膝关节', '[13]'),
    ('过劳性膝痛', '[14]'),
    ('神经肌肉热身', '[15]'),
]

added_count = 0
for keyword, citation in citation_keywords:
    for para in doc.paragraphs:
        if keyword in para.text and citation not in para.text:
            para.add_run(citation)
            added_count += 1
            print(f'已添加 {citation}')
            break

print(f'共添加 {added_count} 个引用')

# 保存
output_path = r'C:\Users\Administrator\Desktop\毕业论文_刘小龙_最终版.docx'
doc.save(output_path)
print(f'保存完成: {output_path}')

# 验证
doc2 = Document(output_path)
all_citations = []
for para in doc2.paragraphs:
    matches = re.findall(r'\[(\d+)\]', para.text)
    all_citations.extend(matches)

print(f'\n验证: 共 {len(all_citations)} 个引用, {len(set(all_citations))} 个唯一引用')
print('引用列表:', sorted(all_citations, key=int))

if len(all_citations) == len(set(all_citations)) == 15:
    print('✓ 引用正确: 15篇文献，无重复')
else:
    print('✗ 引用有问题')
